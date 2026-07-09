"""One-time script to generate the NDA .docx template with Jinja2 tags for docxtpl."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MUTUAL NON-DISCLOSURE AGREEMENT")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# Intro
doc.add_paragraph(
    'This Mutual Non-Disclosure Agreement ("Agreement") is entered into as of '
    "{{ effective_date }} (the \"Effective Date\") by and between:"
)

doc.add_paragraph(
    "{{ disclosing_party }}, a {{ disclosing_entity_type }} organized under the laws of "
    "{{ disclosing_jurisdiction }} (\"Disclosing Party\"); and"
)

doc.add_paragraph(
    "{{ receiving_party }}, a {{ receiving_entity_type }} organized under the laws of "
    "{{ receiving_jurisdiction }} (\"Receiving Party\")."
)

doc.add_paragraph(
    "Each a \"Party\" and collectively, the \"Parties.\""
)

doc.add_paragraph()

# Section 1: Purpose
h = doc.add_paragraph()
run = h.add_run("1. PURPOSE")
run.bold = True
doc.add_paragraph(
    "The Parties wish to explore a potential business relationship concerning "
    "{{ purpose }} (the \"Purpose\"). In connection with the Purpose, each Party "
    "may disclose certain confidential and proprietary information to the other Party."
)

# Section 2: Confidential Information
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("2. DEFINITION OF CONFIDENTIAL INFORMATION")
run.bold = True
doc.add_paragraph(
    "\"Confidential Information\" means any non-public information disclosed by either "
    "Party to the other Party, whether orally, in writing, or in any other form, that "
    "is designated as confidential or that reasonably should be understood to be "
    "confidential given the nature of the information and the circumstances of disclosure."
)

# Section 3: Obligations
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("3. OBLIGATIONS OF RECEIVING PARTY")
run.bold = True
doc.add_paragraph(
    "The Receiving Party shall: (a) hold the Confidential Information in strict confidence; "
    "(b) not disclose the Confidential Information to any third party except as permitted "
    "herein; (c) use the Confidential Information solely for the Purpose; and (d) protect "
    "the Confidential Information using at least the same degree of care it uses to protect "
    "its own confidential information, but no less than reasonable care."
)

# Section 4: Exclusions
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("4. EXCLUSIONS")
run.bold = True
doc.add_paragraph(
    "Confidential Information does not include information that: (a) is or becomes publicly "
    "available through no fault of the Receiving Party; (b) was known to the Receiving Party "
    "prior to disclosure; (c) is independently developed by the Receiving Party without "
    "reference to the Confidential Information; (d) is rightfully received from a third "
    "party without restriction; or (e) is required to be disclosed by law, regulation, or "
    "court order, provided that the Receiving Party gives the Disclosing Party prompt written "
    "notice and cooperates in seeking a protective order."
)

# Section 5: Term
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("5. TERM")
run.bold = True
doc.add_paragraph(
    "This Agreement shall remain in effect for a period of {{ term_months }} months from "
    "the Effective Date, unless terminated earlier by either Party with {{ notice_days }} "
    "days' written notice. The obligations of confidentiality shall survive termination "
    "for a period of {{ survival_years }} years."
)

# Section 6: Return of Information
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("6. RETURN OF INFORMATION")
run.bold = True
doc.add_paragraph(
    "Upon termination of this Agreement or upon request of the Disclosing Party, the "
    "Receiving Party shall promptly return or destroy all Confidential Information and "
    "any copies thereof, except for copies retained in accordance with the Receiving "
    "Party's standard backup and archival procedures or as required by applicable law."
)

# Section 7: Governing Law
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("7. GOVERNING LAW")
run.bold = True
doc.add_paragraph(
    "This Agreement shall be governed by and construed in accordance with the laws of "
    "the State of {{ governing_law }}, without regard to its conflict of laws principles."
)

# Section 8: Remedies
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("8. REMEDIES")
run.bold = True
doc.add_paragraph(
    "Each Party acknowledges that a breach of this Agreement may cause irreparable harm "
    "for which monetary damages may be inadequate. Accordingly, either Party may seek "
    "equitable relief, including injunction and specific performance, in addition to any "
    "other remedies available at law or in equity."
)

# Section 9: General
doc.add_paragraph()
h = doc.add_paragraph()
run = h.add_run("9. GENERAL PROVISIONS")
run.bold = True
doc.add_paragraph(
    "This Agreement constitutes the entire agreement between the Parties with respect to "
    "the subject matter hereof and supersedes all prior negotiations, representations, or "
    "agreements relating thereto. This Agreement may not be amended except in writing signed "
    "by both Parties. Neither Party may assign this Agreement without the prior written "
    "consent of the other Party."
)

# Signature block
doc.add_paragraph()
doc.add_paragraph()

sig = doc.add_paragraph("IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.")

doc.add_paragraph()
doc.add_paragraph("{{ disclosing_party }}")
doc.add_paragraph("By: _________________________________")
doc.add_paragraph("Name: {{ disclosing_signatory }}")
doc.add_paragraph("Title: {{ disclosing_title }}")
doc.add_paragraph("Date: _________________________________")

doc.add_paragraph()
doc.add_paragraph("{{ receiving_party }}")
doc.add_paragraph("By: _________________________________")
doc.add_paragraph("Name: {{ receiving_signatory }}")
doc.add_paragraph("Title: {{ receiving_title }}")
doc.add_paragraph("Date: _________________________________")

doc.save("data/templates/nda_mutual.docx")
print("Created data/templates/nda_mutual.docx")
