"""
Tests for the Cicero template engine integration.

Covers: cicero_bridge module, compliance_draft Cicero path,
        playbook fail-fast, conditional sections, determinism,
        markdown-to-docx conversion, and audit logging.
"""
from __future__ import annotations

import json
import os
import sqlite3
import tempfile
from pathlib import Path

import pytest

import shutil

from contract_drafting.cicero_bridge import (
    CiceroDraftResult,
    draft,
    markdown_to_docx,
    _build_data,
    _compute_data_hash,
    _resolve_template_dir,
    _read_template,
    _read_template_version,
    _homebrew_to_docx,
    _REFERENCE_DOCX,
)

HAS_PANDOC = (
    shutil.which("pandoc") is not None
    or Path("/opt/homebrew/bin/pandoc").is_file()
)
from contract_drafting.compliance_draft import (
    DraftRequest,
    draft_contract,
    get_audit_log,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def basic_request() -> DraftRequest:
    return DraftRequest(
        disclosing_party="TestCo",
        receiving_party="AcmeCorp",
        effective_date="2026-01-15",
    )


@pytest.fixture
def tmp_db(tmp_path) -> str:
    return str(tmp_path / "test.db")


@pytest.fixture
def tmp_docx(tmp_path) -> str:
    return str(tmp_path / "output.docx")


# ---------------------------------------------------------------------------
# TestCiceroBridge — template engine internals
# ---------------------------------------------------------------------------

class TestCiceroBridge:

    def test_resolve_template_dir(self):
        """Template directory for nda-mutual exists."""
        d = _resolve_template_dir("nda-mutual")
        assert d.is_dir()
        assert (d / "text" / "grammar.tem.md").exists()
        assert (d / "package.json").exists()

    def test_resolve_template_dir_missing(self):
        """Nonexistent template raises FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            _resolve_template_dir("nonexistent-template-xyz")

    def test_read_template_version(self):
        """Version is read from package.json."""
        d = _resolve_template_dir("nda-mutual")
        version = _read_template_version(d)
        assert version == "1.0.0"

    def test_build_data_json(self, basic_request):
        """DraftRequest snake_case maps to camelCase template vars."""
        data = _build_data(basic_request)
        assert data["disclosingParty"] == "TestCo"
        assert data["receivingParty"] == "AcmeCorp"
        assert data["effectiveDate"] == "2026-01-15"
        assert data["termMonths"] == 24
        assert data["noticeDays"] == 30
        assert data["survivalYears"] == 3
        assert data["governingLaw"] == "Washington"
        assert data["mutual"] is True
        assert data["hasNonCompete"] is False

    def test_data_hash_deterministic(self, basic_request):
        """Same input produces same SHA256 hash."""
        d1 = _build_data(basic_request)
        d2 = _build_data(basic_request)
        assert _compute_data_hash(d1) == _compute_data_hash(d2)

    def test_data_hash_varies(self):
        """Different input produces different hash."""
        r1 = DraftRequest(disclosing_party="A", receiving_party="B", effective_date="2026-01-01")
        r2 = DraftRequest(disclosing_party="X", receiving_party="Y", effective_date="2026-01-01")
        h1 = _compute_data_hash(_build_data(r1))
        h2 = _compute_data_hash(_build_data(r2))
        assert h1 != h2

    def test_draft_nda_smoke(self, basic_request):
        """End-to-end draft produces text containing party names."""
        result = draft(basic_request)
        assert result.success is True
        assert result.template_name == "nda-mutual"
        assert result.template_version == "1.0.0"
        assert len(result.data_hash) == 64  # SHA256 hex
        assert "TestCo" in result.text
        assert "AcmeCorp" in result.text
        assert "2026-01-15" in result.text
        assert "24 months" in result.text
        assert "Washington" in result.text

    def test_draft_with_non_compete(self):
        """hasNonCompete=True includes non-compete section."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            effective_date="2026-01-01", has_non_compete=True,
        )
        result = draft(req)
        assert result.success
        assert "NON-COMPETITION" in result.text

    def test_draft_without_non_compete(self):
        """hasNonCompete=False excludes non-compete section."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            effective_date="2026-01-01", has_non_compete=False,
        )
        result = draft(req)
        assert result.success
        assert "NON-COMPETITION" not in result.text

    def test_draft_with_non_solicitation(self):
        """hasNonSolicitation=True includes non-solicitation section."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            effective_date="2026-01-01", has_non_solicitation=True,
        )
        result = draft(req)
        assert result.success
        assert "NON-SOLICITATION" in result.text

    def test_draft_with_residuals(self):
        """hasResidualsClause=True includes residuals section."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            effective_date="2026-01-01", has_residuals_clause=True,
        )
        result = draft(req)
        assert result.success
        assert "RESIDUALS" in result.text

    def test_draft_bad_template(self):
        """Nonexistent template dir returns success=False."""
        req = DraftRequest(disclosing_party="A", receiving_party="B")
        result = draft(req, template_name="nonexistent-xyz")
        assert result.success is False
        assert result.error is not None

    def test_draft_deterministic(self):
        """Same request produces identical output."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-01",
        )
        r1 = draft(req)
        r2 = draft(req)
        assert r1.text == r2.text
        assert r1.data_hash == r2.data_hash

    def test_markdown_to_docx(self, tmp_docx):
        """Markdown-to-docx conversion produces a valid .docx file."""
        md = "# Title\n\nSome **bold** text.\n\n## Section\n\nParagraph here.\n\n---\n\nAfter rule."
        path = markdown_to_docx(md, tmp_docx)
        assert Path(path).exists()
        assert Path(path).suffix == ".docx"
        # Verify it's a valid docx (zip with expected structure)
        import zipfile
        assert zipfile.is_zipfile(path)

    def test_markdown_to_docx_content(self, tmp_docx):
        """Markdown-to-docx preserves heading text and bold text (any backend)."""
        from docx import Document
        md = "# Main Title\n\nSome **bold** text and normal text.\n\n## Sub Heading\n\nAnother paragraph."
        path = markdown_to_docx(md, tmp_docx)
        doc = Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Main Title" in full_text
        assert "Sub Heading" in full_text
        assert "bold" in full_text

    def test_homebrew_fallback(self, tmp_docx):
        """Homebrew converter still works as a standalone fallback."""
        md = "# Title\n\nSome **bold** text.\n\n## Section\n\nParagraph."
        path = _homebrew_to_docx(md, tmp_docx)
        assert Path(path).exists()
        import zipfile
        assert zipfile.is_zipfile(path)

    @pytest.mark.skipif(not HAS_PANDOC, reason="pandoc not installed")
    def test_pandoc_subclauses(self, tmp_docx):
        """Pandoc renders (a)(b)(c) sub-clauses correctly."""
        md = (
            "## Section\n\n"
            "The Party shall:\n\n"
            "(a) do the first thing;\n\n"
            "(b) do the second thing;\n\n"
            "(c) do the third thing.\n"
        )
        path = markdown_to_docx(md, tmp_docx)
        from docx import Document
        doc = Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "first thing" in full_text
        assert "second thing" in full_text
        assert "third thing" in full_text

    @pytest.mark.skipif(not HAS_PANDOC, reason="pandoc not installed")
    def test_pandoc_nda_full_roundtrip(self, tmp_docx):
        """Full NDA template renders through pandoc without error."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15",
        )
        result = draft(req)
        assert result.success
        path = markdown_to_docx(result.text, tmp_docx)
        assert Path(path).exists()
        from docx import Document
        doc = Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "TestCo" in full_text
        assert "AcmeCorp" in full_text
        assert "Confidential Information" in full_text

    def test_read_template_version_missing_package_json(self, tmp_path):
        """Missing package.json defaults to version 0.0.0."""
        assert _read_template_version(tmp_path) == "0.0.0"


# ---------------------------------------------------------------------------
# TestDraftPipeline — full compliance_draft integration
# ---------------------------------------------------------------------------

class TestDraftPipeline:

    def test_cicero_pass(self, basic_request, tmp_db, tmp_docx):
        """Standard request passes playbook and generates .docx."""
        basic_request.output_path = tmp_docx
        result = draft_contract(basic_request, engine="cicero", db_path=tmp_db)

        assert result["gate_result"] == "PASS"
        assert result["engine"] == "cicero"
        assert result["mode"] == "draft"
        assert result["template_name"] == "nda-mutual"
        assert result["data_hash"]
        assert Path(result["output_path"]).exists()

    def test_cicero_pass_has_disclaimer(self, basic_request, tmp_db, tmp_docx):
        """A successful Cicero draft carries the single-sourced not-legal-advice footer."""
        from contract_drafting.demo_mars_beat import LEGAL_DISCLAIMER

        basic_request.output_path = tmp_docx
        result = draft_contract(basic_request, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "PASS"
        assert result.get("disclaimer") == LEGAL_DISCLAIMER

    def test_llm_pass_has_disclaimer(self, monkeypatch, tmp_db, tmp_docx):
        """A successful --engine llm draft carries the SAME single-sourced footer
        as the Cicero path (parity; no engine produces a draft without it)."""
        from contract_drafting import compliance_draft as cd
        from contract_drafting.demo_mars_beat import LEGAL_DISCLAIMER

        monkeypatch.setattr(cd, "_call_llm_for_fields", lambda *a, **k: {
            "purpose": "x", "term_months": 24, "notice_days": 30,
            "survival_years": 3, "governing_law": "New York",
        })
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="New York",
            output_path=tmp_docx,
        )
        result = draft_contract(req, engine="llm", db_path=tmp_db)
        assert result["engine"] == "llm"
        assert result["gate_result"] == "PASS"
        assert result.get("disclaimer") == LEGAL_DISCLAIMER

    def test_no_disclaimer_on_non_pass_result(self, tmp_db, tmp_docx):
        """The PASS-specific disclaimer must NOT be in the result dict for non-PASS gates
        (Codex P2): an ESCALATED draft (120-month term) carries no 'PASS certifies...' field,
        and a BLOCKED draft (non-compete) carries none either."""
        esc = draft_contract(DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            term_months=120, effective_date="2026-01-01", output_path=tmp_docx,
        ), engine="cicero", db_path=tmp_db)
        assert esc["gate_result"] == "ESCALATED"
        assert "disclaimer" not in esc
        blk = draft_contract(DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            has_non_compete=True, effective_date="2026-01-01",
        ), engine="cicero", db_path=tmp_db)
        assert blk["gate_result"] == "BLOCKED"
        assert "disclaimer" not in blk

    def test_cicero_blocked_non_compete(self, tmp_db):
        """Non-compete triggers BLOCKED before drafting."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            has_non_compete=True, effective_date="2026-01-01",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)

        assert result["gate_result"] == "BLOCKED"
        assert result["output_path"] is None
        assert any(v["clause_type"] == "Non-Compete" for v in result["violations"])

    def test_cicero_blocked_non_solicitation(self, tmp_db):
        """Non-solicitation triggers BLOCKED."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            has_non_solicitation=True, effective_date="2026-01-01",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"

    def test_cicero_escalated_long_term(self, tmp_db, tmp_docx):
        """120-month term gets ESCALATED but still generates output."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            term_months=120, effective_date="2026-01-01",
            output_path=tmp_docx,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)

        assert result["gate_result"] == "ESCALATED"
        assert result["output_path"] is not None
        assert Path(result["output_path"]).exists()
        assert any(v["clause_type"] == "Term" for v in result["violations"])

    def test_cicero_blocked_empty_disclosing_party(self, tmp_db):
        """Empty disclosing party name returns BLOCKED."""
        req = DraftRequest(
            disclosing_party="", receiving_party="AcmeCorp",
            effective_date="2026-01-01",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert "error" in result
        assert "disclosingParty" in result["error"]

    def test_cicero_blocked_empty_receiving_party(self, tmp_db):
        """Empty receiving party name returns BLOCKED."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="",
            effective_date="2026-01-01",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert "error" in result
        assert "receivingParty" in result["error"]

    def test_cicero_failure_has_mode_and_engine(self, tmp_db):
        """Cicero engine failure returns mode and engine keys for CLI handler."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            doc_type="nonexistent_type",
            effective_date="2026-01-01",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        # Should fail because there's no template for nonexistent_type
        assert result["gate_result"] == "BLOCKED"
        assert result.get("mode") == "draft"
        assert result.get("engine") == "cicero"

    def test_cicero_audit_log(self, basic_request, tmp_db, tmp_docx):
        """Audit log records Cicero template name, data hash, and full input data."""
        basic_request.output_path = tmp_docx
        result = draft_contract(basic_request, engine="cicero", db_path=tmp_db)

        logs = get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert len(logs) == 1
        log_entry = logs[0]
        assert "cicero/nda-mutual" in log_entry["template_id"]
        assert "data_hash" in log_entry["notes"]
        # Verify full input data is stored (not just hash)
        slot_values = json.loads(log_entry["slot_values"])
        assert "disclosingParty" in slot_values
        assert slot_values["disclosingParty"] == "TestCo"
        assert "data_hash" in slot_values

    def test_cicero_integration_full(self, tmp_db, tmp_path):
        """Full integration: .docx produced + audit log + correct gate."""
        out = str(tmp_path / "integration_test.docx")
        req = DraftRequest(
            disclosing_party="AlphaCo",
            receiving_party="BetaCorp",
            term_months=24,
            governing_law="Delaware",
            effective_date="2026-06-01",
            output_path=out,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)

        assert result["gate_result"] == "PASS"
        assert result["engine"] == "cicero"
        assert Path(out).exists()

        # Check audit trail
        logs = get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert len(logs) == 1
        assert logs[0]["gate_result"] == "PASS"
        assert logs[0]["output_path"] == out


# ---------------------------------------------------------------------------
# TestPlaybookValidation — playbook rules for drafting
# ---------------------------------------------------------------------------

class TestPlaybookValidation:

    def test_standard_terms_pass(self, tmp_db, tmp_docx):
        """Standard NDA terms within playbook range pass."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            term_months=24, governing_law="Washington",
            effective_date="2026-01-01", output_path=tmp_docx,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "PASS"
        assert "violations" not in result

    def test_max_standard_term_passes(self, tmp_db, tmp_docx):
        """36-month term (3 years) is at the upper bound and passes."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            term_months=36, effective_date="2026-01-01",
            output_path=tmp_docx,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "PASS"

    def test_over_max_term_escalates(self, tmp_db, tmp_docx):
        """37-month term exceeds playbook standard and escalates."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            term_months=37, effective_date="2026-01-01",
            output_path=tmp_docx,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "ESCALATED"

    def test_invalid_governing_law_blocked(self, tmp_db):
        """Invalid governing law (not a US state) returns BLOCKED."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            governing_law="Wahington",  # typo
            effective_date="2026-01-01",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert "governingLaw" in result["error"]

    def test_invalid_date_format_blocked(self, tmp_db):
        """Invalid date format returns BLOCKED."""
        req = DraftRequest(
            disclosing_party="A", receiving_party="B",
            effective_date="01/15/2026",  # wrong format
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert "effectiveDate" in result["error"]

    def test_cicero_escalates_other_entity_sentinel(self, tmp_db):
        """Codex P2 + beat-6 semantics: entityType=OTHER_ENTITY on the DETERMINISTIC Cicero path
        must fail closed with a clean abstain handoff (parity with governingLaw=OTHER and the LLM
        path) -- not a generic 'Cicero draft failed' from deep in _build_data with no `abstained`
        flag. An abstention is NOT a playbook violation: it routes to ESCALATED (human review),
        renders nothing, and always writes an audit row carrying the raw captured ask."""
        from contract_drafting.compliance_draft import get_audit_log
        req = DraftRequest(
            disclosing_party="A", receiving_party="B", effective_date="2026-01-15",
            receiving_entity_type="OTHER_ENTITY", receiving_entity_type_raw="GmbH",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "ESCALATED"
        assert result.get("abstained") is True
        assert result.get("output_path") is None  # fail-closed: nothing rendered
        assert "OTHER_ENTITY" in result.get("error", "")
        assert "Cicero draft failed" not in result.get("error", "")
        assert result.get("engine") == "cicero"
        # audit row always written, with the abstained field + raw value + record sha
        rows = get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert rows and rows[0]["gate_result"] == "ESCALATED"
        slot = json.loads(rows[0]["slot_values"])
        assert slot["abstentions"] == [
            {"field": "receivingEntityType", "sentinel": "OTHER_ENTITY", "raw": "GmbH"}]
        assert len(slot["escalation_sha256"]) == 64


# ---------------------------------------------------------------------------
# TestSchemaValidator — Concerto-generated JSON Schema validation
# ---------------------------------------------------------------------------

from contract_drafting.schema_validator import validate_template_data, clear_cache


class TestSchemaValidator:

    def _valid_data(self) -> dict:
        return {
            "disclosingParty": "TestCo", "receivingParty": "AcmeCorp",
            "effectiveDate": "2026-01-15", "governingLaw": "Washington",
            "disclosingEntityType": "corporation", "receivingEntityType": "corporation",
            "purpose": "test", "termMonths": 24, "noticeDays": 30, "survivalYears": 3,
            "mutual": True, "hasNonCompete": False, "hasNonSolicitation": False,
            "hasResidualsClause": False,
        }

    def test_valid_data_passes(self):
        errors = validate_template_data(self._valid_data())
        assert errors == []

    def test_empty_party_fails(self):
        data = self._valid_data()
        data["disclosingParty"] = ""
        errors = validate_template_data(data)
        assert any("disclosingParty" in e for e in errors)

    def test_invalid_state_fails(self):
        data = self._valid_data()
        data["governingLaw"] = "Wahington"
        errors = validate_template_data(data)
        assert any("governingLaw" in e for e in errors)

    def test_invalid_date_fails(self):
        data = self._valid_data()
        data["effectiveDate"] = "01/15/2026"
        errors = validate_template_data(data)
        assert any("effectiveDate" in e for e in errors)

    def test_wrong_type_fails(self):
        data = self._valid_data()
        data["termMonths"] = "twenty-four"
        errors = validate_template_data(data)
        assert any("termMonths" in e for e in errors)

    def test_missing_schema(self):
        errors = validate_template_data({}, template_name="nonexistent-xyz")
        assert len(errors) == 1
        assert "not found" in errors[0].lower()

    def test_all_us_states_valid(self):
        # governingLaw is now a native Jurisdiction enum whose values are
        # identifiers; the production path normalizes display names to
        # identifiers (jurisdiction_map.to_identifier) before validating. This
        # test mirrors that pipeline: display name -> identifier -> validate.
        from contract_drafting.jurisdiction_map import to_identifier
        for state in ["Alabama", "New Hampshire", "District of Columbia", "Wyoming"]:
            data = self._valid_data()
            data["governingLaw"] = to_identifier(state)
            errors = validate_template_data(data)
            gov_errors = [e for e in errors if "governingLaw" in e]
            assert not gov_errors, f"{state} should be valid"

    def test_cache_clear(self):
        clear_cache()
        errors = validate_template_data(self._valid_data())
        assert errors == []


# ---------------------------------------------------------------------------
# TestTemplateRegistry — Template auto-discovery
# ---------------------------------------------------------------------------

from contract_drafting.template_registry import TemplateRegistry, get_registry


class TestTemplateRegistry:

    def test_scan_finds_nda(self):
        registry = TemplateRegistry.scan()
        assert "nda-mutual" in registry.templates
        info = registry.get("nda-mutual")
        assert info.version == "1.0.0"
        assert info.has_grammar is True

    def test_get_for_doc_type_nda(self):
        registry = TemplateRegistry.scan()
        info = registry.get_for_doc_type("nda")
        assert info is not None
        assert info.name == "nda-mutual"

    def test_get_for_doc_type_unknown(self):
        registry = TemplateRegistry.scan()
        assert registry.get_for_doc_type("employment-agreement") is None

    def test_scan_empty_dir(self, tmp_path):
        registry = TemplateRegistry.scan(tmp_path)
        assert len(registry.templates) == 0

    def test_list_types(self):
        registry = TemplateRegistry.scan()
        assert "nda-mutual" in registry.list_types()


# ---------------------------------------------------------------------------
# TestVersionTracking — Audit trail version tracking
# ---------------------------------------------------------------------------

class TestVersionTracking:

    def test_audit_records_model_version(self, tmp_db, tmp_path):
        out = str(tmp_path / "version_test.docx")
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", output_path=out,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        logs = get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert len(logs) == 1
        assert logs[0].get("model_version") == "1.0.0"

    def test_audit_records_schema_hash(self, tmp_db, tmp_path):
        out = str(tmp_path / "hash_test.docx")
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", output_path=out,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        logs = get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert len(logs) == 1
        schema_hash = logs[0].get("schema_hash")
        assert schema_hash is not None
        assert len(schema_hash) == 64


# ---------------------------------------------------------------------------
# TestAnchorScanning — {{> partial}} syntax parsing
# ---------------------------------------------------------------------------

from contract_drafting.cicero_bridge import (  # noqa: E402
    _scan_anchors,
    _load_partials,
    _clear_partials_cache,
    draft_with_data,
)


class TestAnchorScanning:
    """Verify _scan_anchors extracts {{> name}} partial references correctly."""

    def test_no_anchors(self):
        assert _scan_anchors("plain text with {{variable}} but no partials") == set()

    def test_single_anchor(self):
        assert _scan_anchors("{{> governing-law}}") == {"governing-law"}

    def test_multiple_anchors(self):
        text = """
        # Contract
        {{> governing-law}}
        {{> dispute-resolution}}
        {{> signature-block}}
        """
        assert _scan_anchors(text) == {
            "governing-law",
            "dispute-resolution",
            "signature-block",
        }

    def test_duplicate_anchors_deduplicated(self):
        text = "{{> foo}} some text {{> foo}}"
        assert _scan_anchors(text) == {"foo"}

    def test_whitespace_variants(self):
        assert _scan_anchors("{{>foo}}") == {"foo"}
        assert _scan_anchors("{{> foo }}") == {"foo"}
        assert _scan_anchors("{{>  foo  }}") == {"foo"}

    def test_ignores_regular_variables(self):
        assert _scan_anchors("{{partyAName}} and {{partyBName}}") == set()

    def test_ignores_conditionals(self):
        assert _scan_anchors("{{#if hasFoo}}body{{/if}}") == set()

    def test_mixed_content(self):
        text = "Party {{partyAName}} signs. {{> signature-block}} {{#if mutual}}yes{{/if}}"
        assert _scan_anchors(text) == {"signature-block"}


class TestAnchorScanningProperties:
    """Property-based tests for the anchor regex — fuzz with Hypothesis."""

    def test_roundtrip_valid_anchor_names(self):
        """Any valid anchor name wrapped in {{> X}} should roundtrip."""
        from hypothesis import given, strategies as st

        @given(st.from_regex(r"^[a-z][a-z0-9-]{0,30}$", fullmatch=True))
        def check(name: str):
            template = f"prefix {{{{> {name}}}}} suffix"
            assert _scan_anchors(template) == {name}

        check()

    def test_multiple_random_anchors(self):
        """Any set of random valid anchor names should all be extracted."""
        from hypothesis import given, strategies as st

        @given(
            st.sets(
                st.from_regex(r"^[a-z][a-z0-9-]{0,20}$", fullmatch=True),
                min_size=1,
                max_size=8,
            )
        )
        def check(names: set[str]):
            template = "\n".join(f"{{{{> {n}}}}}" for n in names)
            assert _scan_anchors(template) == names

        check()


class TestPartialLoading:
    """Verify _load_partials cache and error handling."""

    def setup_method(self):
        _clear_partials_cache()

    def test_empty_required_returns_empty(self, tmp_path):
        assert _load_partials(tmp_path, set()) == {}

    def test_missing_shared_dir_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="Shared clauses directory"):
            _load_partials(tmp_path / "nonexistent", {"foo"})

    def test_loads_existing_partials(self, tmp_path):
        (tmp_path / "foo.md").write_text("foo content")
        (tmp_path / "bar.md").write_text("bar content")
        result = _load_partials(tmp_path, {"foo", "bar"})
        assert result == {"foo": "foo content", "bar": "bar content"}

    def test_missing_required_partial_raises(self, tmp_path):
        (tmp_path / "foo.md").write_text("foo content")
        with pytest.raises(FileNotFoundError, match="missing shared clauses.*ghost"):
            _load_partials(tmp_path, {"foo", "ghost"})

    def test_only_returns_required_subset(self, tmp_path):
        """Even if more partials exist on disk, only required ones are returned."""
        (tmp_path / "wanted.md").write_text("yes")
        (tmp_path / "unwanted.md").write_text("no")
        result = _load_partials(tmp_path, {"wanted"})
        assert result == {"wanted": "yes"}

    def test_cache_hit_avoids_reread(self, tmp_path):
        """Second call with same mtime hits cache."""
        f = tmp_path / "foo.md"
        f.write_text("first")
        _load_partials(tmp_path, {"foo"})
        # Mutate file directly without touching mtime — cache should still hold first version
        original_mtime = f.stat().st_mtime
        f.write_text("second")
        os.utime(f, (original_mtime, original_mtime))  # reset mtime
        result = _load_partials(tmp_path, {"foo"})
        assert result == {"foo": "first"}, "cache should have held original content"

    def test_cache_invalidates_on_mtime_change(self, tmp_path):
        """New mtime → re-read from disk."""
        import time
        f = tmp_path / "foo.md"
        f.write_text("v1")
        _load_partials(tmp_path, {"foo"})
        time.sleep(0.01)
        f.write_text("v2")
        # mtime naturally advances; no utime reset
        result = _load_partials(tmp_path, {"foo"})
        assert result == {"foo": "v2"}, "cache should have invalidated"


class TestCompositionEngine:
    """End-to-end rendering with the new anchor-based composition."""

    def test_jv_regression_single_signature(self):
        """IRON RULE: the v0.4.0.1 JV duplicate-signature bug cannot return."""
        data = _jv_test_data()
        result = draft_with_data(data, template_name="joint-venture")
        assert result.success, f"draft failed: {result.error}"
        sig_count = result.text.count("Signature: _________________________") // 2
        assert sig_count == 1, f"duplicate signature regressed: {sig_count} blocks"

    def test_jv_governing_law_before_signature(self):
        """IRON RULE: governing-law clause must appear BEFORE signature block."""
        data = _jv_test_data()
        result = draft_with_data(data, template_name="joint-venture")
        gov_pos = result.text.find("Applicable Law")
        sig_pos = result.text.find("IN WITNESS WHEREOF")
        assert gov_pos != -1 and sig_pos != -1
        assert gov_pos < sig_pos, "governing-law must render before signature"

    def test_anchors_resolved_no_placeholders(self):
        """No unrendered {{> X}} anchors in the output."""
        data = _jv_test_data()
        result = draft_with_data(data, template_name="joint-venture")
        import re
        assert not re.search(r"\{\{>\s*\w", result.text), \
            "unrendered partial anchor in output"

    def test_missing_partial_raises_clear_error(self, tmp_path, monkeypatch):
        """If a template references a partial that doesn't exist, draft fails loudly."""
        # Mirror the real layout: templates/cicero/<name> + templates/shared-clauses
        templates_base = tmp_path / "templates" / "cicero"
        templates_base.mkdir(parents=True)
        tpl_root = templates_base / "broken"
        (tpl_root / "text").mkdir(parents=True)
        (tpl_root / "text" / "grammar.tem.md").write_text(
            "# Broken\n\n{{> does-not-exist}}\n"
        )
        (tpl_root / "package.json").write_text('{"version": "0.0.1"}')
        # Shared-clauses exists but is empty (so the "missing required" path fires)
        (tmp_path / "templates" / "shared-clauses").mkdir()

        # Point _TEMPLATES_BASE at our fixture
        from contract_drafting import cicero_bridge
        monkeypatch.setattr(cicero_bridge, "_TEMPLATES_BASE", templates_base)
        _clear_partials_cache()

        result = draft_with_data({}, template_name="broken")
        assert not result.success
        assert "missing shared clauses" in result.error
        assert "does-not-exist" in result.error

    def test_variant_flags_merged(self):
        """composition.json variants end up in the render data dict."""
        data = _jv_test_data()
        result = draft_with_data(data, template_name="joint-venture")
        # JV has useNegotiationFirst=true → dispute-resolution clause shows "friendly negotiation"
        assert "friendly negotiation" in result.text

    def test_all_four_templates_render(self):
        """All 4 migrated templates render without anchor errors."""
        for name in ["joint-venture", "strategic-cooperation", "intermediary", "partnership"]:
            data = _jv_test_data()  # over-provisioned dict, extras ignored
            result = draft_with_data(data, template_name=name)
            assert result.success, f"{name} failed: {result.error}"
            assert len(result.text) > 1000, f"{name} render suspiciously short"

    def test_determinism(self):
        """Same input → same output hash."""
        data = _jv_test_data()
        r1 = draft_with_data(data, template_name="joint-venture")
        r2 = draft_with_data(data, template_name="joint-venture")
        assert r1.data_hash == r2.data_hash
        assert r1.text == r2.text


class TestGoldenFiles:
    """Verify agent E2E output matches locked-in golden snapshots.

    If a template or shared clause changes, these tests fail. Regenerate
    the snapshots DELIBERATELY by running:
        venv/bin/python -c "from tests.regenerate_golden import main; main()"
    """

    GOLDEN_DIR = Path(__file__).parent / "fixtures" / "golden"

    @pytest.mark.parametrize("name,template,data_fn", [
        ("agent_e2e_joint_venture", "joint-venture", "_golden_jv_v2"),
        ("agent_e2e_intermediary", "intermediary", "_golden_intermediary"),
        ("agent_e2e_strategic_cooperation", "strategic-cooperation", "_golden_cooperation"),
        ("agent_e2e_partnership", "partnership", "_golden_partnership"),
        ("agent_e2e_consulting", "consulting", "_golden_consulting"),
    ])
    def test_golden_match(self, name, template, data_fn):
        data = globals()[data_fn]()
        result = draft_with_data(data, template_name=template)
        assert result.success, f"{name} render failed: {result.error}"

        golden_path = self.GOLDEN_DIR / f"{name}.md"
        if not golden_path.exists():
            pytest.skip(f"golden file not found: {golden_path}")

        golden = golden_path.read_text(encoding="utf-8")
        assert result.text == golden, (
            f"{name} output drift from golden snapshot. "
            f"If intentional, regenerate the snapshot."
        )


def _golden_jv() -> dict:
    return {
        "partyAName": "Acme AI Inc.", "partyBName": "Beacon Trading Corporation",
        "partyAAddress": "1-6-1 Marunouchi, Chiyoda-ku, Tokyo",
        "partyBAddress": "67 Sejong-daero, Jung-gu, Seoul",
        "partyALegalRep": "Jamie Tan, President",
        "partyBLegalRep": "Chris Lee, Managing Director",
        "partyAContact": "jamie.tan@example.com", "partyBContact": "chris.lee@example.com",
        "jvCompanyName": "Acme-Beacon APAC Holdings Pte. Ltd.",
        "registeredCapital": "SGD 10,000,000",
        "registrationPlace": "Singapore",
        "operationTermYears": 15, "partyAEquityPercent": 50, "partyBEquityPercent": 50,
        "contributionDeadlineDays": 30, "operatingRegion": "Southeast Asia",
        "profitDistributionDays": 60,
        "applicableCompanyLaw": "Singapore Companies Act",
        "totalDirectors": 7, "partyADirectors": 3, "partyBDirectors": 3,
        "chairmanNominatedBy": "Acme AI Inc.",
        "confidentialityPenaltyAmount": "SGD 1,000,000",
        "breachCurePeriodDays": 45,
        "governingLaw": "Republic of Singapore", "arbitrationBody": "SIAC",
        "effectiveDate": "2026-05-01",
    }


def _golden_intermediary() -> dict:
    return {
        "agreementNo": "DEMO-0001", "signingPlace": "Singapore",
        "partyAName": "Acme AI Inc.",
        "partyALegalRep": "Alex Kim, CEO",
        "partyAAddress": "1 Raffles Quay, Level 33, #33-01, Singapore 048583",
        "partyAContact": "alex.kim@example.com",
        "partyBName": "Riverbend Tech Partners",
        "partyBLegalRep": "Dana Pham, Managing Director",
        "partyBAddress": "235 Dong Khoi Street, District 1, Ho Chi Minh City, Vietnam",
        "partyBContact": "+84 28 3822 5678",
        "partyBRoleDescription": "Vietnamese market intermediary with strong connections to data center operators, major telecommunications operators (Viettel, VNPT), and fintech companies",
        "targetCountry": "Vietnam",
        "finderFeePercent": 5, "finderFeeCap": "no cap",
        "paymentTermDays": 45, "confidentialityTermYears": 3,
        "terminationNoticeDays": 90, "breachCurePeriodDays": 30,
        "exclusivityExpiryMonths": 6, "agreementTermMonths": 36,
        "governingLaw": "Republic of Singapore",
        "arbitrationBody": "Singapore International Arbitration Centre (SIAC)",
        "effectiveDate": "2026-05-15",
    }


def _golden_cooperation() -> dict:
    return {
        "agreementNo": "DEMO-0003", "signingPlace": "Singapore",
        "partyAName": "Acme AI Inc.",
        "partyALegalRep": "Alex Kim, CEO",
        "partyAAddress": "1 Raffles Quay, #33-01, Singapore 048583",
        "partyAContact": "alex.kim@example.com",
        "partyBName": "Horizon Tech Nigeria Ltd",
        "partyBLegalRep": "Adebayo Okafor, Managing Director",
        "partyBAddress": "14 Adeyemi Lawson Close, Ikoyi, Lagos",
        "partyBContact": "adebayo@horizon-tech.ng",
        "agreementTermYears": 5, "terminationNoticeDays": 90,
        "confidentialityTermYears": 3,
        "governingLaw": "Republic of Singapore", "arbitrationBody": "SIAC",
        "effectiveDate": "2026-05-10",
    }


class TestTemplateLinter:
    """Wrap validate_templates CLI in a pytest so CI catches broken anchors."""

    def test_all_template_anchors_resolve(self):
        from contract_drafting.validate_templates import validate_all
        exit_code, messages = validate_all()
        assert exit_code == 0, f"Broken anchors found: {messages}"


def _golden_partnership() -> dict:
    return {
        "partyAName": "Acme AI Inc.",
        "partyBName": "Savannah Innovation Kenya Ltd",
        "partyAAddress": "1 Raffles Quay, #33-01, Singapore 048583",
        "partyBAddress": "ABC Place, Waiyaki Way, Nairobi",
        "partyALegalRep": "Alex Kim, CEO",
        "partyBLegalRep": "Wanjiru Kamau, CEO",
        "partyAContact": "alex.kim@example.com",
        "partyBContact": "wanjiru@savannah.co.ke",
        "targetCountry": "Kenya", "operatingRegion": "East Africa",
        "partyAEquityPercent": 60, "partyBEquityPercent": 40,
        "agreementTermYears": 10, "confidentialityTermYears": 5,
        "dataProtectionLaw": "Kenya Data Protection Act 2019",
        "expenditureThresholdUSD": "100,000",
        "contractThresholdUSD": "500,000",
        "settlementThresholdUSD": "50,000",
        "governingLaw": "Kenya",
        "arbitrationBody": "Nairobi Centre for International Arbitration (NCIA)",
        "effectiveDate": "2026-06-01",
    }


def _jv_test_data() -> dict:
    """Over-provisioned fixture with every camelCase field any of our templates needs."""
    return {
        # Party info (shared across all 4 templates)
        "partyAName": "Acme AI Inc.",
        "partyBName": "Beacon Trading Corporation",
        "partyAAddress": "Tokyo",
        "partyBAddress": "Seoul",
        "partyALegalRep": "Alice",
        "partyBLegalRep": "Bob",
        "partyAContact": "a@test",
        "partyBContact": "b@test",
        "partyBRoleDescription": "APAC market partner",
        # JV-specific
        "jvCompanyName": "Acme-Beacon APAC Holdings",
        "registeredCapital": "SGD 10,000,000",
        "registrationPlace": "Singapore",
        "operationTermYears": 15,
        "partyAEquityPercent": 50,
        "partyBEquityPercent": 50,
        "contributionDeadlineDays": 30,
        "operatingRegion": "Southeast Asia",
        "profitDistributionDays": 60,
        "applicableCompanyLaw": "Singapore Companies Act",
        "totalDirectors": 7,
        "partyADirectors": 3,
        "partyBDirectors": 3,
        "chairmanNominatedBy": "Acme AI Inc.",
        "confidentialityPenaltyAmount": "SGD 1,000,000",
        # Intermediary-specific
        "agreementNo": "TEST-001",
        "signingPlace": "Singapore",
        "targetCountry": "Vietnam",
        "finderFeePercent": 5,
        "finderFeeCap": "no cap",
        "paymentTermDays": 45,
        "exclusivityExpiryMonths": 6,
        "agreementTermMonths": 36,
        # Strategic-cooperation / partnership shared
        "agreementTermYears": 3,
        "renewalTermYears": 3,
        "terminationNoticeDays": 90,
        "confidentialityTermYears": 3,
        "dataProtectionLaw": "PDPA",
        "expenditureThresholdUSD": "100,000",
        "contractThresholdUSD": "500,000",
        "settlementThresholdUSD": "50,000",
        # Common tail
        "breachCurePeriodDays": 30,
        "governingLaw": "Republic of Singapore",
        "arbitrationBody": "SIAC",
        "effectiveDate": "2026-05-15",
    }


# ---------------------------------------------------------------------------
# TestClauseRefinement — refine.py primitives
# ---------------------------------------------------------------------------

from contract_drafting import refine as refine_mod  # noqa: E402
from contract_drafting.refine import (  # noqa: E402
    RefinedClause,
    ClauseDiff,
    ValidationResult,
    refine_clause,
    diff_clause,
    validate_refined_clause,
    promote_clause,
    list_overlays,
    _extract_variables,
    _extract_conditional_opens,
    _strip_code_fences,
    OVERLAYS_DIR,
)


class TestClauseRefinement:
    """Unit tests for refine_clause with LLM mocked."""

    def _mock_llm_response(self, monkeypatch, response: str) -> list[tuple[str, str]]:
        """Install a mock _call_refine_llm that returns `response` and records all calls."""
        calls: list[tuple[str, str]] = []
        def fake(system_prompt: str, user_prompt: str) -> str:
            calls.append((system_prompt, user_prompt))
            return response
        monkeypatch.setattr(refine_mod, "_call_refine_llm", fake)
        return calls

    def test_happy_path_rewrite(self, monkeypatch):
        # Must preserve both {{#useFormalGoverningLaw}} and {{^useFormalGoverningLaw}}
        # conditionals present in the canonical clause.
        self._mock_llm_response(monkeypatch,
            "## Governing Law\n\n"
            "{{#useFormalGoverningLaw}}\n"
            "**Applicable Law:** The laws of {{governingLaw}} shall govern this Agreement.\n"
            "{{/useFormalGoverningLaw}}\n"
            "{{^useFormalGoverningLaw}}\n"
            "**Applicable Law:** {{governingLaw}} law applies.\n"
            "{{/useFormalGoverningLaw}}\n"
        )
        data = {"governingLaw": "Singapore"}
        r = refine_clause("governing-law", "make phrasing slightly different",
                         data=data, template_name="joint-venture")
        assert r.status == "PASS", r.reason
        assert r.validation.ok
        assert "Applicable Law" in r.after
        assert r.diff.changed
        assert r.sha256_before and r.sha256_after

    def test_allowed_vars_enforced(self, monkeypatch):
        # LLM returns a clause that references an unknown variable.
        self._mock_llm_response(monkeypatch,
            "**Applicable Law:** Laws of {{unknownJurisdiction}} apply."
        )
        data = {"governingLaw": "Singapore"}
        r = refine_clause("governing-law", "use a new jurisdiction var",
                         data=data, template_name="joint-venture")
        assert r.status == "BLOCKED"
        assert "unknownJurisdiction" in r.validation.unknown_vars

    def test_conditional_preservation(self, monkeypatch):
        # Rewrite dispute-resolution but drop the {{#useNegotiationFirst}} section
        self._mock_llm_response(monkeypatch,
            "**Dispute Resolution:** Submit to {{arbitrationBody}} for arbitration."
        )
        data = {"arbitrationBody": "SIAC"}
        r = refine_clause("dispute-resolution", "simplify",
                         data=data, template_name="joint-venture")
        assert r.status == "BLOCKED"
        assert "useNegotiationFirst" in r.validation.dropped_conditionals

    def test_mustache_syntax_error(self, monkeypatch):
        # LLM returns unclosed section
        self._mock_llm_response(monkeypatch, "{{#open}} no close tag here")
        data = {"governingLaw": "Singapore"}
        r = refine_clause("governing-law", "break it",
                         data=data, template_name="joint-venture")
        assert r.status == "BLOCKED"
        assert any("Mustache" in e or "mustache" in e.lower()
                   for e in r.validation.errors)

    def test_needs_new_var_escape_hatch(self, monkeypatch):
        self._mock_llm_response(monkeypatch,
            "NEEDS_NEW_VAR: ipCarveoutScope\n(would need this variable)"
        )
        data = {"governingLaw": "Singapore"}
        r = refine_clause("governing-law", "add IP carveout",
                         data=data, template_name="joint-venture")
        assert r.status == "BLOCKED"
        assert "NEEDS_NEW_VAR" in r.reason
        assert "ipCarveoutScope" in r.reason

    def test_playbook_violation_escape_hatch(self, monkeypatch):
        self._mock_llm_response(monkeypatch,
            "PLAYBOOK_VIOLATION: non-solicitation clauses are prohibited"
        )
        data = {"governingLaw": "Singapore"}
        r = refine_clause("governing-law", "add non-solicit",
                         data=data, template_name="joint-venture")
        assert r.status == "BLOCKED"
        assert "PLAYBOOK_VIOLATION" in r.reason

    def test_strips_code_fences(self, monkeypatch):
        # LLM ignores the rule and wraps in code fences anyway. Preserve conditionals.
        self._mock_llm_response(monkeypatch,
            "```\n## Governing Law\n\n"
            "{{#useFormalGoverningLaw}}Formal: {{governingLaw}}.{{/useFormalGoverningLaw}}\n"
            "{{^useFormalGoverningLaw}}Simple: {{governingLaw}}.{{/useFormalGoverningLaw}}\n```"
        )
        data = {"governingLaw": "Singapore"}
        r = refine_clause("governing-law", "simplify",
                         data=data, template_name="joint-venture")
        assert r.status == "PASS", r.reason
        assert "```" not in r.after

    def test_refine_does_not_mutate_disk(self, monkeypatch):
        """Canonical shared-clauses/ files must not change during refinement."""
        from contract_drafting.refine import SHARED_CLAUSES_DIR
        before_content = (SHARED_CLAUSES_DIR / "governing-law.md").read_text()
        self._mock_llm_response(monkeypatch,
            "**Applicable Law:** Laws of {{governingLaw}} apply."
        )
        data = {"governingLaw": "Singapore"}
        refine_clause("governing-law", "rewrite",
                      data=data, template_name="joint-venture")
        after_content = (SHARED_CLAUSES_DIR / "governing-law.md").read_text()
        assert before_content == after_content

    def test_refine_unknown_clause_raises(self, monkeypatch):
        self._mock_llm_response(monkeypatch, "anything")
        with pytest.raises(FileNotFoundError):
            refine_clause("nonexistent-clause", "any instruction",
                          data={}, template_name="joint-venture")

    def test_audit_logs_refinement_event(self, monkeypatch, tmp_path):
        db = str(tmp_path / "audit.db")
        # Preserve both conditionals so status=PASS is asserted below.
        self._mock_llm_response(monkeypatch,
            "## Governing Law\n"
            "{{#useFormalGoverningLaw}}F: {{governingLaw}}.{{/useFormalGoverningLaw}}\n"
            "{{^useFormalGoverningLaw}}S: {{governingLaw}}.{{/useFormalGoverningLaw}}\n"
        )
        data = {"governingLaw": "Singapore"}
        refine_clause("governing-law", "test audit",
                      data=data, template_name="joint-venture",
                      db_path=db, user="test-agent")
        # Read the audit log
        import sqlite3
        conn = sqlite3.connect(db)
        rows = conn.execute(
            "SELECT doc_type, gate_result, notes FROM audit_log"
        ).fetchall()
        conn.close()
        assert len(rows) == 1
        doc_type, gate_result, notes = rows[0]
        assert doc_type == "refinement"
        assert gate_result == "PASS"
        payload = json.loads(notes)
        assert payload["event"] == "clause_refinement"
        assert payload["clause_name"] == "governing-law"
        assert payload["instruction"] == "test audit"
        assert payload["sha256_before"] and payload["sha256_after"]

    def test_idempotent_with_same_mock_response(self, monkeypatch):
        self._mock_llm_response(monkeypatch,
            "## Governing Law\n"
            "{{#useFormalGoverningLaw}}F: {{governingLaw}}.{{/useFormalGoverningLaw}}\n"
            "{{^useFormalGoverningLaw}}S: {{governingLaw}}.{{/useFormalGoverningLaw}}\n"
        )
        data = {"governingLaw": "Singapore"}
        r1 = refine_clause("governing-law", "same instruction",
                           data=data, template_name="joint-venture")
        r2 = refine_clause("governing-law", "same instruction",
                           data=data, template_name="joint-venture")
        assert r1.sha256_after == r2.sha256_after
        assert r1.after == r2.after


# ---------------------------------------------------------------------------
# TestPromoteClause — promote_clause + overlay filesystem
# ---------------------------------------------------------------------------

class TestPromoteClause:

    def _cleanup_overlay(self, clause_name: str, overlay_id: str) -> None:
        overlay_path = OVERLAYS_DIR / clause_name / f"{overlay_id}.md"
        meta_path = OVERLAYS_DIR / clause_name / f"{overlay_id}.meta.json"
        for p in (overlay_path, meta_path):
            if p.exists():
                p.unlink()

    def test_promote_writes_overlay_and_meta(self):
        self._cleanup_overlay("governing-law", "test-promote-1")
        try:
            p = promote_clause(
                overlay_id="test-promote-1",
                clause_name="governing-law",
                text="**Applicable Law:** Delaware.",
                meta={"source_draft_hash": "abc123"},
            )
            assert Path(p.overlay_path).is_file()
            assert Path(p.meta_path).is_file()
            meta = json.loads(Path(p.meta_path).read_text())
            assert meta["overlay_id"] == "test-promote-1"
            assert meta["clause_name"] == "governing-law"
            assert meta["source_draft_hash"] == "abc123"
            assert "sha256" in meta and "created_at" in meta
            assert "playbook_version" in meta
        finally:
            self._cleanup_overlay("governing-law", "test-promote-1")

    def test_promote_collision_raises(self):
        self._cleanup_overlay("governing-law", "test-collide")
        try:
            promote_clause("test-collide", "governing-law", "first")
            with pytest.raises(FileExistsError, match="already exists"):
                promote_clause("test-collide", "governing-law", "second")
        finally:
            self._cleanup_overlay("governing-law", "test-collide")

    def test_promote_unknown_clause_rejected(self):
        with pytest.raises(FileNotFoundError, match="unknown clause"):
            promote_clause("any-id", "nonexistent-clause", "text")

    def test_promote_bad_overlay_id_rejected(self):
        with pytest.raises(ValueError, match="Invalid overlay_id"):
            promote_clause("UPPERCASE-BAD", "governing-law", "text")
        with pytest.raises(ValueError):
            promote_clause("has spaces", "governing-law", "text")
        with pytest.raises(ValueError):
            promote_clause("123startswith-digit", "governing-law", "text")

    def test_list_overlays_shows_promoted(self):
        self._cleanup_overlay("governing-law", "test-list-1")
        self._cleanup_overlay("governing-law", "test-list-2")
        try:
            promote_clause("test-list-1", "governing-law", "a")
            promote_clause("test-list-2", "governing-law", "b")
            overlays = list_overlays("governing-law")
            assert "governing-law" in overlays
            assert "test-list-1" in overlays["governing-law"]
            assert "test-list-2" in overlays["governing-law"]
        finally:
            self._cleanup_overlay("governing-law", "test-list-1")
            self._cleanup_overlay("governing-law", "test-list-2")

    def test_promoted_overlay_used_in_render(self):
        """overlay_map={'governing-law': 'id'} causes render to use overlay text."""
        self._cleanup_overlay("governing-law", "test-render-overlay")
        try:
            overlay_text = "**Applicable Law:** The laws of OVERLAY-TEST apply."
            promote_clause("test-render-overlay", "governing-law", overlay_text)
            r = draft_with_data(
                _jv_test_data(), template_name="joint-venture",
                overlay_map={"governing-law": "test-render-overlay"},
            )
            assert r.success, r.error
            assert "OVERLAY-TEST" in r.text
        finally:
            self._cleanup_overlay("governing-law", "test-render-overlay")

    def test_promote_audit_logged(self, tmp_path):
        self._cleanup_overlay("governing-law", "test-audit-promote")
        db = str(tmp_path / "audit.db")
        try:
            promote_clause(
                "test-audit-promote", "governing-law", "text",
                db_path=db,
                meta={"source_draft_hash": "deadbeef"},
            )
            conn = sqlite3.connect(db)
            rows = conn.execute(
                "SELECT doc_type, gate_result, notes FROM audit_log"
            ).fetchall()
            conn.close()
            assert len(rows) == 1
            doc_type, gate_result, notes = rows[0]
            assert doc_type == "promotion"
            assert gate_result == "PASS"
            payload = json.loads(notes)
            assert payload["event"] == "clause_promotion"
            assert payload["overlay_id"] == "test-audit-promote"
        finally:
            self._cleanup_overlay("governing-law", "test-audit-promote")


# ---------------------------------------------------------------------------
# TestDiffAndValidate — standalone tests for helpers
# ---------------------------------------------------------------------------

class TestDiffAndValidate:

    def test_diff_unchanged(self):
        d = diff_clause("same", "same")
        assert d.changed is False
        assert d.added_words == []
        assert d.removed_words == []

    def test_diff_structured_output(self):
        d = diff_clause("the quick brown fox", "the slow red fox")
        assert d.changed is True
        assert "slow" in d.added_words
        assert "red" in d.added_words
        assert "quick" in d.removed_words
        assert "brown" in d.removed_words

    def test_validate_ok_with_allowed_vars(self):
        v = validate_refined_clause(
            "Use {{foo}} and {{bar}}",
            allowed_vars={"foo", "bar", "baz"},
            original_text="",
        )
        assert v.ok

    def test_validate_rejects_unknown_var(self):
        v = validate_refined_clause(
            "Use {{ghost}}",
            allowed_vars={"foo"},
            original_text="",
        )
        assert not v.ok
        assert "ghost" in v.unknown_vars

    def test_validate_preserves_conditional(self):
        v = validate_refined_clause(
            "Plain text no conditional",
            allowed_vars=set(),
            original_text="{{#flag}}something{{/flag}}",
        )
        assert not v.ok
        assert "flag" in v.dropped_conditionals

    def test_extract_variables_skips_partials_and_conditionals(self):
        assert _extract_variables("{{> partial}}") == set()
        assert _extract_variables("{{#cond}}x{{/cond}}") == set()
        assert _extract_variables("{{^not}}x{{/not}}") == set()
        assert _extract_variables("{{real}}") == {"real"}


# ---------------------------------------------------------------------------
# TestOverlayComposition — extend render path with clause_overrides + overlay_map
# ---------------------------------------------------------------------------

class TestOverlayComposition:

    def test_clause_overrides_ephemeral(self):
        """clause_overrides doesn't touch disk."""
        from contract_drafting.refine import SHARED_CLAUSES_DIR
        canonical_before = (SHARED_CLAUSES_DIR / "governing-law.md").read_text()
        r = draft_with_data(
            _jv_test_data(), template_name="joint-venture",
            clause_overrides={"governing-law":
                              "**Applicable Law:** EPHEMERAL TEST text."},
        )
        assert r.success
        assert "EPHEMERAL TEST" in r.text
        canonical_after = (SHARED_CLAUSES_DIR / "governing-law.md").read_text()
        assert canonical_before == canonical_after

    def test_override_beats_overlay(self):
        """clause_overrides takes precedence over overlay_map."""
        # Setup: promote an overlay
        from contract_drafting.refine import OVERLAYS_DIR
        overlay_id = "test-precedence-x"
        overlay_path = OVERLAYS_DIR / "governing-law" / f"{overlay_id}.md"
        meta_path = OVERLAYS_DIR / "governing-law" / f"{overlay_id}.meta.json"
        try:
            promote_clause(overlay_id, "governing-law",
                           "**Applicable Law:** OVERLAY wins.")
            r = draft_with_data(
                _jv_test_data(), template_name="joint-venture",
                overlay_map={"governing-law": overlay_id},
                clause_overrides={"governing-law":
                                  "**Applicable Law:** OVERRIDE wins."},
            )
            assert r.success
            assert "OVERRIDE wins" in r.text
            assert "OVERLAY wins" not in r.text
        finally:
            for p in (overlay_path, meta_path):
                if p.exists():
                    p.unlink()

    def test_data_hash_differs_with_overlay(self):
        from contract_drafting.refine import OVERLAYS_DIR
        overlay_id = "test-hash-diff"
        overlay_path = OVERLAYS_DIR / "governing-law" / f"{overlay_id}.md"
        meta_path = OVERLAYS_DIR / "governing-law" / f"{overlay_id}.meta.json"
        try:
            promote_clause(overlay_id, "governing-law",
                           "**Applicable Law:** Delaware only.")
            data = _jv_test_data()
            r_canonical = draft_with_data(data, template_name="joint-venture")
            r_overlay = draft_with_data(
                data, template_name="joint-venture",
                overlay_map={"governing-law": overlay_id},
            )
            assert r_canonical.success and r_overlay.success
            assert r_canonical.data_hash != r_overlay.data_hash
        finally:
            for p in (overlay_path, meta_path):
                if p.exists():
                    p.unlink()

    def test_data_hash_differs_with_override(self):
        data = _jv_test_data()
        r1 = draft_with_data(data, template_name="joint-venture")
        r2 = draft_with_data(
            data, template_name="joint-venture",
            clause_overrides={"governing-law": "anything different"},
        )
        assert r1.data_hash != r2.data_hash

    def test_missing_overlay_raises(self):
        r = draft_with_data(
            _jv_test_data(), template_name="joint-venture",
            overlay_map={"governing-law": "does-not-exist"},
        )
        assert not r.success
        assert "Overlay not found" in r.error


# ---------------------------------------------------------------------------
# TestLinterDuplicateAnchor — rejects templates with duplicate {{> X}}
# ---------------------------------------------------------------------------

class TestLinterDuplicateAnchor:

    def test_duplicate_anchor_rejected(self, tmp_path, monkeypatch):
        """A template with the same {{> X}} twice must fail the linter."""
        from contract_drafting import validate_templates, cicero_bridge

        # Build a fixture template tree
        templates_base = tmp_path / "templates" / "cicero"
        templates_base.mkdir(parents=True)
        tpl = templates_base / "dup-test"
        (tpl / "text").mkdir(parents=True)
        (tpl / "text" / "grammar.tem.md").write_text(
            "# Test\n\n{{> governing-law}}\n\nLater:\n\n{{> governing-law}}\n"
        )
        (tpl / "package.json").write_text('{"version": "0.0.1"}')
        shared = tmp_path / "templates" / "shared-clauses"
        shared.mkdir()
        (shared / "governing-law.md").write_text("canonical text")

        monkeypatch.setattr(cicero_bridge, "_TEMPLATES_BASE", templates_base)
        monkeypatch.setattr(validate_templates, "_TEMPLATES_BASE", templates_base)
        cicero_bridge._clear_partials_cache()

        exit_code, messages = validate_templates.validate_all()
        assert exit_code == 1
        assert any("duplicate anchor" in m for m in messages), messages

    def test_real_templates_no_duplicates(self):
        """All 5 real templates must be clean. Would detect a regression."""
        from contract_drafting.validate_templates import validate_all
        exit_code, messages = validate_all()
        assert exit_code == 0, messages


# ---------------------------------------------------------------------------
# TestLinterEnumDisplayMap — RT8: a template declaring @Display enums (or
# enum-backed abstain fields) MUST ship a covering enum-displays.map.json,
# because to_display_enum is fail-open (verbatim pass-through) without it.
# ---------------------------------------------------------------------------

class TestLinterEnumDisplayMap:

    @staticmethod
    def _make_template(tmp_path, *, with_map: bool, stale: bool = False,
                       map_content: dict | None = None) -> Path:
        """A minimal template tree whose model.cto declares one @Display enum
        (EntityType) and one plain enum (which needs no map entry)."""
        templates_base = tmp_path / "templates" / "cicero"
        tpl = templates_base / "enum-test"
        (tpl / "text").mkdir(parents=True)
        (tpl / "text" / "grammar.tem.md").write_text("# Test\n\nNo anchors.\n")
        (tpl / "model").mkdir()
        (tpl / "model" / "model.cto").write_text(
            "namespace test\n\n"
            "enum EntityType {\n"
            '  @Display("corporation")\n'
            "  o corporation\n"
            '  @Display("trust")\n'
            "  o trust\n"
            "}\n\n"
            "enum PlainEnum {\n  o A\n  o B\n}\n"
        )
        (tpl / "package.json").write_text('{"version": "0.0.1"}')
        (tmp_path / "templates" / "shared-clauses").mkdir()
        if with_map:
            if map_content is not None:
                content = map_content
            else:
                content = {} if stale else {
                    "EntityType": {"corporation": "corporation", "trust": "trust"}}
            (tpl / "enum-displays.map.json").write_text(json.dumps(content))
        return templates_base

    def _lint(self, tmp_path, monkeypatch, **kw):
        from contract_drafting import validate_templates
        base = self._make_template(tmp_path, **kw)
        monkeypatch.setattr(validate_templates, "_TEMPLATES_BASE", base)
        return validate_templates.validate_all()

    def test_missing_map_reported(self, tmp_path, monkeypatch):
        """The focused RT8 unit test: a @Display-bearing template WITHOUT the map
        must fail the linter with a loud, named error."""
        exit_code, messages = self._lint(tmp_path, monkeypatch, with_map=False)
        assert exit_code == 1
        assert any("enum-displays.map.json is MISSING" in m for m in messages), messages
        assert any("EntityType" in m for m in messages), messages

    def test_stale_map_reported(self, tmp_path, monkeypatch):
        exit_code, messages = self._lint(tmp_path, monkeypatch,
                                         with_map=True, stale=True)
        assert exit_code == 1
        assert any("STALE" in m and "EntityType" in m for m in messages), messages

    def test_covering_map_passes(self, tmp_path, monkeypatch):
        exit_code, messages = self._lint(tmp_path, monkeypatch, with_map=True)
        assert exit_code == 0, messages

    def test_member_drift_missing_member_reported(self, tmp_path, monkeypatch):
        """C11: the lint is MEMBER-level -- an enum member declared in the .cto but
        absent from the map (here: trust) is a lint error, not a silent fail-open."""
        exit_code, messages = self._lint(
            tmp_path, monkeypatch, with_map=True,
            map_content={"EntityType": {"corporation": "corporation"}})
        assert exit_code == 1
        assert any("STALE" in m and "trust" in m and "missing member" in m
                   for m in messages), messages

    def test_member_drift_stale_member_reported(self, tmp_path, monkeypatch):
        """C11: a map key that is no longer an enum member (here: ghost) is a stale
        codegen artifact and a lint error."""
        exit_code, messages = self._lint(
            tmp_path, monkeypatch, with_map=True,
            map_content={"EntityType": {"corporation": "corporation",
                                        "trust": "trust", "ghost": "ghost"}})
        assert exit_code == 1
        assert any("STALE" in m and "ghost" in m and "no longer" in m
                   for m in messages), messages

    def test_plain_enum_without_display_needs_no_map(self, tmp_path, monkeypatch):
        """A template whose .cto has only un-decorated enums (and no abstain policy)
        must not be forced to ship a map."""
        from contract_drafting import validate_templates
        base = tmp_path / "templates" / "cicero"
        tpl = base / "plain-test"
        (tpl / "text").mkdir(parents=True)
        (tpl / "text" / "grammar.tem.md").write_text("# Test\n\nNo anchors.\n")
        (tpl / "model").mkdir()
        (tpl / "model" / "model.cto").write_text(
            "namespace test\n\nenum PlainEnum {\n  o A\n  o B\n}\n")
        (tmp_path / "templates" / "shared-clauses").mkdir()
        monkeypatch.setattr(validate_templates, "_TEMPLATES_BASE", base)
        exit_code, messages = validate_templates.validate_all()
        assert exit_code == 0, messages

    def test_shipped_templates_pass_the_enum_display_lint(self):
        """The real templates (incl. nda-mutual with Jurisdiction/EntityType/
        DisputeForum) must satisfy the new lint."""
        from contract_drafting.validate_templates import validate_all
        exit_code, messages = validate_all()
        assert exit_code == 0, messages


# ---------------------------------------------------------------------------
# v0.7.0.0 — Golden template absorption: Phase 1 (new shared clauses)
# ---------------------------------------------------------------------------

from contract_drafting.cicero_bridge import _normalize_array_fields  # noqa: E402
import chevron  # noqa: E402

_SHARED_DIR = Path(__file__).resolve().parent.parent / "data" / "templates" / "shared-clauses"


def _render_partial(name: str, data: dict) -> str:
    """Render a single shared-clause file with chevron, after array normalization."""
    text = (_SHARED_DIR / f"{name}.md").read_text(encoding="utf-8")
    normalized = _normalize_array_fields(data)
    return chevron.render(text, normalized)


class TestPhase1NewSharedClauses:
    """v0.7.0.0 Phase 1 — three new shared clauses + array normalization."""

    def test_normalize_empty_list(self):
        out = _normalize_array_fields({"localComplianceLaws": []})
        assert out["hasLocalComplianceLaws"] is False
        assert out["localComplianceLawsJoined"] == ""
        assert out["localComplianceLaws"] == []

    def test_normalize_single_item_list(self):
        out = _normalize_array_fields({"localComplianceLaws": ["FCPA"]})
        assert out["hasLocalComplianceLaws"] is True
        assert out["localComplianceLawsJoined"] == "FCPA"

    def test_normalize_multi_item_list(self):
        out = _normalize_array_fields({
            "localComplianceLaws": ["FCPA", "UK Bribery Act", "Kenyan law"]
        })
        assert out["hasLocalComplianceLaws"] is True
        assert out["localComplianceLawsJoined"] == "FCPA, UK Bribery Act, Kenyan law"

    def test_normalize_non_list_passthrough(self):
        out = _normalize_array_fields({"name": "Acme", "count": 3})
        assert "hasName" not in out
        assert "nameJoined" not in out
        assert out["name"] == "Acme"

    def test_normalize_idempotent(self):
        d1 = _normalize_array_fields({"localComplianceLaws": ["FCPA"]})
        d2 = _normalize_array_fields(d1)
        assert d1 == d2

    def test_normalize_user_supplied_has_field_wins(self):
        """If caller explicitly sets `hasFoo`, derived value must NOT override.

        Use case: caller has a non-empty list but wants to suppress the
        conditional section in this specific render. Explicit user intent wins.
        """
        out = _normalize_array_fields({
            "localComplianceLaws": ["FCPA"],
            "hasLocalComplianceLaws": False,  # explicit override
        })
        assert out["hasLocalComplianceLaws"] is False
        assert out["localComplianceLawsJoined"] == "FCPA"

    def test_normalize_skips_non_string_lists(self):
        # Lists of dicts (e.g. for Mustache section iteration) untouched.
        out = _normalize_array_fields({"items": [{"k": "v"}]})
        assert "hasItems" not in out
        assert "itemsJoined" not in out

    def test_representations_warranties_renders(self):
        out = _render_partial("representations-warranties", {"governingLaw": "Republic of Singapore"})
        # Content-only partial: heading comes from the including template's Roman-numeral section.
        assert "severally and independently represents, warrants, and undertakes" in out
        assert "Republic of Singapore" in out
        assert "{{" not in out, "unrendered placeholder remains"

    def test_compliance_renders_empty_local_laws(self):
        out = _render_partial("compliance", {"localComplianceLaws": []})
        # No leading-comma garbage when array is empty
        assert "(ITAR), and other applicable" in out
        # Make sure the conditional section text didn't leak
        assert "{{" not in out

    def test_compliance_renders_single_law(self):
        out = _render_partial("compliance", {"localComplianceLaws": ["Kenyan law"]})
        assert "(ITAR), Kenyan law, and other applicable" in out
        assert "{{" not in out

    def test_compliance_renders_multi_law(self):
        out = _render_partial("compliance", {
            "localComplianceLaws": ["FCPA", "UK Bribery Act", "Kenyan law"]
        })
        assert "(ITAR), FCPA, UK Bribery Act, Kenyan law, and other applicable" in out
        assert "{{" not in out

    def test_compliance_independent_liability_text_present(self):
        out = _render_partial("compliance", {"localComplianceLaws": []})
        assert "**Independent Liability.**" in out
        assert "**Right to Suspend.**" in out

    def test_notices_renders_with_change_days(self):
        out = _render_partial("notices", {"noticeChangeBusinessDays": 5})
        # Content-only partial: heading comes from the including template's Roman-numeral section.
        assert "**Notices and Delivery.**" in out
        assert "within 5 business days" in out
        assert "{{" not in out

    def test_anti_corruption_with_local_laws(self):
        out = _render_partial("anti-corruption", {
            "localComplianceLaws": ["the laws of Republic of Kenya", "Hong Kong AML laws"],
        })
        assert "FCPA" in out
        assert "the laws of Republic of Kenya, Hong Kong AML laws" in out
        assert "{{" not in out

    def test_anti_corruption_backward_compat_no_laws(self):
        # Old template that doesn't pass localComplianceLaws still renders cleanly.
        out = _render_partial("anti-corruption", {})
        assert "FCPA" in out
        assert "UK Bribery Act 2010." in out
        # The {{#hasLocalComplianceLaws}} block is empty → no extra text
        assert "and ." not in out
        assert "{{" not in out

    def test_anti_corruption_empty_array_renders_clean(self):
        out = _render_partial("anti-corruption", {"localComplianceLaws": []})
        assert "UK Bribery Act 2010." in out
        # Empty array → no "and ," fragment leaks
        assert ", and ." not in out

    def test_existing_templates_still_render_after_anti_corruption_change(self):
        """REGRESSION: anti-corruption change must not break existing 4 templates."""
        for name in ["joint-venture", "strategic-cooperation", "intermediary", "partnership"]:
            data = _jv_test_data()
            result = draft_with_data(data, template_name=name)
            assert result.success, f"{name} regressed: {result.error}"
            # If anti-corruption is included in the template, FCPA appears
            if "{{> anti-corruption}}" in (
                _resolve_template_dir(name) / "text" / "grammar.tem.md"
            ).read_text():
                assert "FCPA" in result.text


# ---------------------------------------------------------------------------
# v0.7.0.0 — Golden template absorption: Phase 2 (consulting template)
# ---------------------------------------------------------------------------


def _golden_consulting() -> dict:
    return {
        "agreementNo": "DEMO-0002",
        "signingPlace": "Singapore",
        "partyAName": "Acme AI Inc.",
        "partyAAddress": "1 Example Way, #01-01 Example Tower, Singapore 000000",
        "partyAContact": "+65 0000 0000",
        "partyALegalRep": "Alex Kim, CEO",
        "partyBName": "Savannah Consulting Kenya Ltd",
        "partyBAddress": "ABC Place, Waiyaki Way, Nairobi",
        "partyBLegalRep": "Wanjiru Kamau",
        "partyBContact": "+254 20 4444 5555",
        "partyBRoleDescription": "Kenya market consultant with regulatory and government relations expertise",
        "targetCountry": "Republic of Kenya",
        "consultingFeePercent": 5.0,
        "consultingFeeCap": "USD 250,000",
        "paymentTermDays": 30,
        "confidentialityTermYears": 3,
        "agreementTermMonths": 12,
        "terminationNoticeDays": 30,
        "breachCurePeriodDays": 15,
        "tailPeriodMonths": 6,
        "localComplianceLaws": ["the laws of Republic of Kenya", "the laws of Hong Kong"],
        "governingLaw": "Republic of Singapore",
        "arbitrationBody": "Singapore International Arbitration Centre (SIAC)",
        "noticeChangeBusinessDays": 5,
        "effectiveDate": "2026-05-15",
    }


class TestPhase2ConsultingTemplate:
    """v0.7.0.0 Phase 2 — new consulting contract type."""

    def test_consulting_renders_baseline(self):
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        assert result.success, f"consulting render failed: {result.error}"
        assert len(result.text) > 1000

    def test_consulting_one_signature_block(self):
        """IRON RULE: the v0.4.0.1 duplicate-signature regression cannot return.

        Exactly 2 occurrences of "Signature: ___..." — one per party (A and B),
        in a single signature block. Any other count is a regression.
        """
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        sig_lines = result.text.count("Signature: _________________________")
        assert sig_lines == 2, (
            f"expected exactly 2 'Signature:' lines (one per party in one block), "
            f"got {sig_lines} — duplicate or partial signature block regressed"
        )

    def test_consulting_one_governing_law_clause(self):
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        # "Applicable Law:" appears once (in the governing-law shared partial)
        assert result.text.count("**Applicable Law:**") == 1

    def test_consulting_tail_period_clause_present(self):
        """The tail-period clause is the marquee counsel-vetted feature.

        Beyond keyword presence, also verify the substantive mechanic: the
        tail period must reference the Consulting Fee survival, the months,
        and the Qualifying Transaction concept that scopes it.
        """
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        assert "Tail Period for Qualifying Transactions" in result.text
        assert "6 months after the expiration or termination" in result.text
        assert "Party B shall remain entitled to the Consulting Fee" in result.text

    def test_consulting_non_exclusive_recital_present(self):
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        assert "non-exclusive" in result.text.lower()

    def test_consulting_no_unsubstituted_placeholders(self):
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        import re
        # Allow signature underscores; reject any {{...}} or {{> ...}} remnants
        assert not re.search(r"\{\{[^}]+\}\}", result.text), \
            "unrendered placeholder in output"

    def test_consulting_registered_in_doc_type_map(self):
        from contract_drafting.template_registry import get_registry
        reg = get_registry(refresh=True)
        info = reg.get_for_doc_type("consulting")
        assert info is not None
        assert info.name == "consulting"

    def test_consulting_fee_cap_omitted_renders_empty(self):
        """When caller omits consultingFeeCap, the renderer treats it as blank.

        Mustache does not pull defaults from the Concerto model — defaults live
        in JSON Schema and are visible to callers, but are NOT auto-applied at
        render time. This test pins that behavior so a refactor doesn't silently
        change it. If we ever wire schema-default population into the bridge,
        this test will fail and force an explicit decision.
        """
        data = _golden_consulting()
        data.pop("consultingFeeCap", None)
        result = draft_with_data(data, template_name="consulting")
        assert result.success
        # Mustache leaves missing vars as empty string. The fee-cap clause becomes:
        # "up to a maximum of  (the "Consulting Fee")." — note the double space.
        assert 'up to a maximum of  (the "Consulting Fee")' in result.text, (
            "fee-cap default ('no cap') unexpectedly populated by render layer; "
            "if schema-default population was added, update this test."
        )

    def test_consulting_local_compliance_laws_render(self):
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        # localComplianceLaws should flow into both anti-corruption and compliance
        assert "the laws of Republic of Kenya" in result.text

    def test_consulting_party_b_role_renders(self):
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        assert "Kenya market consultant" in result.text

    def test_consulting_schema_validates_required_fields(self):
        """Schema marks partyBName as required."""
        from contract_drafting.schema_validator import validate_template_data
        data = _golden_consulting()
        del data["partyBName"]
        errors = validate_template_data(data, template_name="consulting")
        assert errors, "missing partyBName should fail validation"
        assert any("partyBName" in e for e in errors)

    def test_consulting_schema_validates_array_field(self):
        """localComplianceLaws optional array field accepted by schema."""
        from contract_drafting.schema_validator import validate_template_data
        data = _golden_consulting()
        errors = validate_template_data(data, template_name="consulting")
        assert not errors, f"valid data rejected: {errors}"
        # Empty array also valid
        data["localComplianceLaws"] = []
        errors = validate_template_data(data, template_name="consulting")
        assert not errors, f"empty list rejected: {errors}"

    def test_consulting_e2e_smoke(self, tmp_docx):
        """[→E2E] data → render → docx round-trip.

        Beyond file existence and size, crack open the docx and assert that the
        marquee content (party B name, tail-period clause, localComplianceLaws)
        actually survived the pandoc round-trip. A pandoc filter that silently
        stripped Mustache-generated text would otherwise pass.
        """
        result = draft_with_data(_golden_consulting(), template_name="consulting")
        assert result.success
        markdown_to_docx(result.text, tmp_docx)
        out = Path(tmp_docx)
        assert out.is_file()
        assert out.stat().st_size > 5_000

        # Crack open the docx and verify content survives the round-trip.
        from docx import Document
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Savannah Consulting Kenya Ltd" in full_text, "party B name missing in docx"
        assert "Tail Period" in full_text, "tail-period clause missing in docx"
        assert "Republic of Kenya" in full_text, "localComplianceLaws content missing in docx"


# ---------------------------------------------------------------------------
# v0.7.0.0 — Golden template absorption: Phase 3 (JV upgrade)
# ---------------------------------------------------------------------------


def _golden_jv_v2() -> dict:
    """JV golden data with v0.7 new fields."""
    return {
        **_golden_jv(),
        "partyAContributionAmount": "SGD 5,000,000",
        "partyBContributionAmount": "SGD 5,000,000",
        "firstPaymentPercent": 20,
        "firstPaymentTrigger": "the issuance of the business license",
        "noticeChangeBusinessDays": 5,
        "localComplianceLaws": ["the laws of Singapore", "the laws of Japan"],
    }


class TestPhase3JointVentureUpgrade:
    """v0.7.0.0 Phase 3 — JV template gets R&W + Compliance + Notices sections."""

    def test_jv_has_reps_and_warranties_section(self):
        """Section VII is R&W. Strict heading match — section position matters."""
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        assert result.success
        assert "## VII. Representations and Warranties" in result.text
        # And the partial body actually rendered, not just the heading
        assert "severally and independently represents, warrants, and undertakes" in result.text

    def test_jv_has_compliance_section(self):
        """Section VIII is Compliance. Strict heading match — section position matters."""
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        assert "## VIII. Compliance" in result.text
        assert "**General Compliance.**" in result.text

    def test_jv_has_dedicated_notices_section(self):
        """Section XII is Notices. Strict heading match — section position matters."""
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        assert "## XII. Notices" in result.text
        assert "**Notices and Delivery.**" in result.text

    def test_jv_first_payment_percent_renders(self):
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        # Section II uses {{firstPaymentPercent}}
        assert "20%" in result.text or "20 %" in result.text

    def test_jv_party_a_contribution_amount_renders(self):
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        assert "SGD 5,000,000" in result.text

    def test_jv_first_payment_trigger_renders(self):
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        assert "the issuance of the business license" in result.text

    def test_jv_one_signature_block_after_upgrade(self):
        """REGRESSION: still passes after Phase 3 section additions.

        Exactly 2 'Signature:' lines (one per party in one block). Any other
        count is the v0.4.0.1 duplicate-signature regression.
        """
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        sig_lines = result.text.count("Signature: _________________________")
        assert sig_lines == 2, (
            f"expected exactly 2 'Signature:' lines (one per party in one block), "
            f"got {sig_lines} — duplicate or partial signature block regressed"
        )

    def test_jv_one_governing_law_clause_after_upgrade(self):
        """REGRESSION: still passes after section additions."""
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        assert result.text.count("**Applicable Law:**") == 1

    def test_jv_section_numbers_in_order(self):
        """All Roman-numeral section headers appear in sequence."""
        result = draft_with_data(_golden_jv_v2(), template_name="joint-venture")
        # Find positions of each section heading
        sections = ["## I.", "## II.", "## III.", "## IV.", "## V.", "## VI.",
                    "## VII.", "## VIII.", "## IX.", "## X.", "## XI."]
        positions = []
        for s in sections:
            pos = result.text.find(s)
            if pos != -1:
                positions.append((s, pos))
        # Ensure we have at least 8 sections (we add 3 to existing ~8 → 11+)
        assert len(positions) >= 8, f"too few sections: {[s for s, _ in positions]}"
        # Ensure they appear in order
        for i in range(1, len(positions)):
            assert positions[i][1] > positions[i - 1][1], \
                f"section {positions[i][0]} appears before {positions[i - 1][0]}"


# ---------------------------------------------------------------------------
# TestJurisdictionEnum — native Jurisdiction enum + @Display map (Lane A: B1/C)
# ---------------------------------------------------------------------------

from contract_drafting import jurisdiction_map as _jmap

_NDA_DIR = Path(__file__).resolve().parent.parent / "data" / "templates" / "cicero" / "nda-mutual"


def _nda_valid_data() -> dict:
    return {
        "disclosingParty": "TestCo", "receivingParty": "AcmeCorp",
        "effectiveDate": "2026-01-15", "governingLaw": "Washington",
        "disclosingEntityType": "corporation", "receivingEntityType": "corporation",
        "purpose": "test", "termMonths": 24, "noticeDays": 30, "survivalYears": 3,
        "mutual": True, "hasNonCompete": False, "hasNonSolicitation": False,
        "hasResidualsClause": False,
    }


class TestJurisdictionEnum:

    @pytest.mark.parametrize("display", [
        "Alabama", "Wyoming",                          # single-word: identifier == display
        "New Hampshire", "District of Columbia",       # multi-word US
        "Republic of Singapore", "England and Wales",  # international
        "People's Republic of China",                  # apostrophe exception
    ])
    def test_display_normalizes_and_validates(self, display):
        data = _nda_valid_data()
        data["governingLaw"] = _jmap.to_identifier(display)
        errors = [e for e in validate_template_data(data) if "governingLaw" in e]
        assert errors == [], f"{display!r} should normalize to a valid identifier"

    def test_mars_rejected_via_enum(self):
        data = _nda_valid_data()
        data["governingLaw"] = _jmap.to_identifier("Mars")  # pass-through, unknown
        errors = [e for e in validate_template_data(data) if "governingLaw" in e]
        assert errors and "allowed values" in errors[0]

    def test_to_display_fail_closed(self):
        with pytest.raises(ValueError):
            _jmap.to_display("Totally_Unknown_Place")

    def test_no_map_template_is_identity(self):
        # Templates without a native Jurisdiction enum (no jurisdictions.map.json)
        # must be identity no-ops, NOT crash. Regression: the migration added an
        # unconditional to_identifier() in _draft_cicero that raised
        # FileNotFoundError for consulting/jv/etc.
        _jmap.clear_cache()
        assert _jmap.to_identifier("California", template_name="consulting") == "California"
        assert _jmap.to_display("Anything", template_name="consulting") == "Anything"
        assert _jmap.to_identifier("Republic of Singapore", template_name="joint-venture") == "Republic of Singapore"
        # nda-mutual (which HAS a map) stays fail-closed for unknown identifiers.
        with pytest.raises(ValueError):
            _jmap.to_display("Totally_Unknown_Place", template_name="nda-mutual")

    def test_nonNDA_cicero_draft_does_not_crash(self, tmp_path):
        # Regression: a non-NDA deterministic draft must not raise
        # FileNotFoundError on the missing jurisdictions.map.json. It may BLOCK on
        # schema/playbook, but must return a dict, never crash.
        req = DraftRequest(
            doc_type="consulting", disclosing_party="A", receiving_party="B",
            effective_date="2026-01-15", governing_law="California",
        )
        result = draft_contract(req, engine="cicero", db_path=str(tmp_path / "t.db"))
        assert isinstance(result, dict)

    def test_draft_with_data_maps_identifier_not_leaked(self):
        # P1 (Codex gate): the raw-data render path (draft_with_data) bypasses the
        # request pipeline, so it must ALSO map a schema-valid identifier
        # ("New_York") to the display name — never render the underscore form.
        from contract_drafting.cicero_bridge import draft_with_data
        data = {
            "disclosingParty": "TestCo", "receivingParty": "AcmeCorp",
            "effectiveDate": "2026-01-15", "purpose": "x",
            "termMonths": 24, "noticeDays": 30, "survivalYears": 3,
            "governingLaw": "New_York", "mutual": True,
        }
        result = draft_with_data(data, template_name="nda-mutual")
        assert result.success, result.error
        assert "New York" in result.text
        assert "New_York" not in result.text

    def test_identifier_roundtrip_all(self):
        _jmap.clear_cache()
        forward = json.loads((_NDA_DIR / "jurisdictions.map.json").read_text())
        assert len(forward) == 65  # 64 jurisdictions + the OTHER abstain sentinel (PR2)
        assert "OTHER" in forward
        assert len(set(forward.values())) == 65, "display names must be unique"
        for identifier, display in forward.items():
            if identifier == "OTHER":
                continue  # abstain sentinel: in the map but to_display refuses to render it
            assert _jmap.to_display(identifier) == display
            assert _jmap.to_identifier(display) == identifier

    def test_map_keys_match_schema_enum(self):
        # Display-map and schema enum derive from the SAME .cto enum; guard drift.
        forward = json.loads((_NDA_DIR / "jurisdictions.map.json").read_text())
        schema = json.loads((_NDA_DIR / "schema.json").read_text())
        enum_vals = None
        for defn in schema.get("definitions", {}).values():
            if isinstance(defn, dict) and "enum" in defn and "Washington" in defn["enum"]:
                enum_vals = set(defn["enum"])
                break
        assert enum_vals is not None, "Jurisdiction enum not found in schema.json"
        assert set(forward.keys()) == enum_vals

    def test_render_uses_display_not_identifier(self):
        # identifier in -> human display out, NEVER the underscore identifier.
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="New_York",
        )
        result = draft(req, template_name="nda-mutual")
        assert result.success, result.error
        assert "New York" in result.text
        assert "New_York" not in result.text

    def test_render_entity_type_display_not_identifier(self):
        """Cross-field: now that entityType is a typed enum, a non-default multi-word
        form renders as its display name, never the underscore identifier -- whether the
        caller passes the identifier or the display name."""
        for value in ("limited_liability_company", "limited liability company"):
            req = DraftRequest(
                disclosing_party="TestCo", receiving_party="AcmeCorp",
                effective_date="2026-01-15",
                disclosing_entity_type=value, receiving_entity_type="trust",
            )
            result = draft(req, template_name="nda-mutual")
            assert result.success, result.error
            assert "limited liability company" in result.text
            assert "limited_liability_company" not in result.text
            assert "a trust" in result.text

    def test_dispute_forum_not_auto_rendered(self):
        """Codex P2: disputeForum is a CAPTURED field, NOT auto-inserted into the rendered
        clause -- rendering a speculatively-filled forum would silently flip a material term
        (court vs arbitration) without user intent. The court-default clause always renders,
        even when disputeForum is set; entityType (which IS rendered) still fails closed on an
        un-representable value."""
        from contract_drafting.cicero_bridge import draft_with_data
        base = {
            "disclosingParty": "A", "receivingParty": "B", "effectiveDate": "2026-01-15",
            "disclosingEntityType": "corporation", "receivingEntityType": "corporation",
            "purpose": "x", "termMonths": 24, "noticeDays": 30, "survivalYears": 3,
            "governingLaw": "New_York", "mutual": True, "hasNonCompete": False,
            "hasNonSolicitation": False, "hasResidualsClause": False,
        }
        # A provided forum does NOT change the venue: the court clause renders, no arbitration.
        with_forum = draft_with_data({**base, "disputeForum": "SIAC"}, template_name="nda-mutual")
        assert with_forum.success, with_forum.error
        assert "courts located in New York" in with_forum.text
        assert "arbitration administered" not in with_forum.text
        # ...but a captured forum is NOT silently dropped: it is surfaced as a warning for human
        # review (Codex P2 -- the "captured for review" promise must be real, not a silent skip).
        assert with_forum.warnings and any("disputeForum" in w or "forum" in w for w in with_forum.warnings)
        # The OTHER_FORUM abstain sentinel surfaces the raw requested forum verbatim for resolution.
        other_forum = draft_with_data(
            {**base, "disputeForum": "OTHER_FORUM", "disputeForumRaw": "ad-hoc UNCITRAL arbitration"},
            template_name="nda-mutual")
        assert other_forum.success and "courts located in New York" in other_forum.text
        assert any("OTHER_FORUM" in w and "ad-hoc UNCITRAL arbitration" in w for w in other_forum.warnings)
        # No forum requested -> no spurious warning (normal NDA drafts stay clean).
        no_forum = draft_with_data(base, template_name="nda-mutual")
        assert no_forum.success and "courts located in New York" in no_forum.text
        assert no_forum.warnings == []
        # entityType IS rendered, so an un-representable form still fails closed (no silent sub).
        bad_entity = draft_with_data({**base, "receivingEntityType": "GmbH"}, template_name="nda-mutual")
        assert not bad_entity.success

    def test_render_international_no_state_wording(self):
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="Republic of Singapore",
        )
        result = draft(req, template_name="nda-mutual")
        assert result.success, result.error
        assert "Republic of Singapore" in result.text
        assert "Republic_of_Singapore" not in result.text
        assert "State of Republic of Singapore" not in result.text  # wording fix


class TestMarsBeatDemo:
    """The constrained-vs-free split-screen (matched scaffolding, one schema)."""

    def test_free_rejected_constrained_valid(self, monkeypatch):
        from contract_drafting import demo_mars_beat as demo

        free_json = json.dumps({**_nda_valid_data(), "governingLaw": "Mars"})
        monkeypatch.setattr(demo, "call_llm", lambda *a, **k: free_json)

        captured = {}
        def fake_structured(question, context, json_schema, **k):
            captured["schema"] = json_schema
            return {**_nda_valid_data(), "governingLaw": "New_York"}
        monkeypatch.setattr(demo, "call_llm_structured", fake_structured)

        result = demo.run_mars_beat("draft an NDA governed by the laws of Mars")
        # Free arm: emits "Mars" -> validator rejects (post-hoc).
        assert result["free"]["valid"] is False
        assert any("governingLaw" in e for e in result["free"]["errors"])
        # Constrained arm: in-enum by construction -> valid.
        assert result["constrained"]["valid"] is True
        # The constraint is actually wired: the bound schema carries the enum.
        sch = json.dumps(captured["schema"])
        assert "Washington" in sch and "New_York" in sch


class TestLegacyLLMEnumConsistency:
    """T10: the --engine llm path now enforces the same Jurisdiction enum."""

    def test_llm_engine_blocks_invalid_jurisdiction(self, monkeypatch, tmp_db):
        from contract_drafting import compliance_draft as cd

        monkeypatch.setattr(cd, "_call_llm_for_fields", lambda *a, **k: {
            "purpose": "x", "term_months": 24, "notice_days": 30,
            "survival_years": 3, "governing_law": "Mars",
        })
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="Mars",
        )
        result = draft_contract(req, engine="llm", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert "governingLaw" in result.get("error", "")

    def test_llm_engine_entity_display_normalized(self, monkeypatch, tmp_db, tmp_docx):
        """Codex P2: the legacy llm path normalizes entity-type display->identifier (parity
        with the cicero path), so a non-default display form validates + renders, not BLOCKED."""
        from contract_drafting import compliance_draft as cd
        monkeypatch.setattr(cd, "_call_llm_for_fields", lambda *a, **k: {
            "purpose": "x", "term_months": 24, "notice_days": 30,
            "survival_years": 3, "governing_law": "New York",
        })
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="New York",
            disclosing_entity_type="limited liability company",  # display form
            output_path=tmp_docx,
        )
        result = draft_contract(req, engine="llm", db_path=tmp_db)
        assert result["gate_result"] == "PASS", result.get("error")

    def test_llm_engine_escalates_other_entity_sentinel(self, monkeypatch, tmp_db, tmp_docx):
        """Codex P2 + beat-6 semantics: the OTHER_ENTITY abstain sentinel must NEVER PASS-render
        on the llm path (parity with governingLaw=OTHER) -- it fails closed (nothing rendered)
        but routes to ESCALATED (human review, not a playbook violation) with an audit row."""
        from contract_drafting import compliance_draft as cd
        monkeypatch.setattr(cd, "_call_llm_for_fields", lambda *a, **k: {
            "purpose": "x", "term_months": 24, "notice_days": 30,
            "survival_years": 3, "governing_law": "New York",
        })
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="New York",
            receiving_entity_type="OTHER_ENTITY", output_path=tmp_docx,
        )
        result = draft_contract(req, engine="llm", db_path=tmp_db)
        assert result["gate_result"] == "ESCALATED"
        assert result.get("abstained") is True
        assert result.get("output_path") is None  # fail-closed: nothing rendered
        assert "OTHER_ENTITY" in result.get("error", "")
        rows = cd.get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert rows and rows[0]["gate_result"] == "ESCALATED"

    def test_llm_engine_nonNDA_not_blocked_by_nda_enum(self, monkeypatch, tmp_db):
        # F3: --engine llm for a NON-NDA doc type must NOT be validated against the
        # nda-mutual Jurisdiction enum. Any error must come from elsewhere (e.g. a
        # missing docx template), never a governingLaw enum rejection.
        from contract_drafting import compliance_draft as cd
        monkeypatch.setattr(cd, "_call_llm_for_fields", lambda *a, **k: {
            "purpose": "x", "term_months": 24, "notice_days": 30,
            "survival_years": 3, "governing_law": "Mars",
        })
        req = DraftRequest(
            doc_type="jv", disclosing_party="A", receiving_party="B",
            effective_date="2026-01-15", governing_law="Mars",
        )
        result = draft_contract(req, engine="llm", db_path=tmp_db)
        err = result.get("error", "")
        assert "allowed values" not in err
        assert "governingLaw" not in err


# ---------------------------------------------------------------------------
# TestAbstainableCodegen — @Abstainable single-sources the abstain instruction
# from model.cto into abstain-policy.json (provenance + drift-guard, NOT a
# behavior change). The composed system prompt MUST stay byte-identical or the
# Gauntlet replay key (which hashes the system prompt) cache-misses.
# ---------------------------------------------------------------------------

# The regen test shells out to scripts/concerto-helper.js, which requires the
# @accordproject/concerto-* packages. node being on PATH is not enough on a fresh
# checkout where `npm install` has not run, so also require node_modules present.
_REPO_ROOT = Path(__file__).resolve().parent.parent
HAS_NODE = (
    shutil.which("node") is not None
    and (_REPO_ROOT / "node_modules" / "@accordproject" / "concerto-core").is_dir()
)

# The EXACT composed abstain system prompt as of the @Abstainable refactor.
# Frozen on purpose: gauntlet.py builds the record/replay cache key by sha256
# over the system prompt, so a single-byte drift here breaks offline replay of
# every constrained_hatch + instr_only-ablation cassette (paper reproducibility).
# If you intend to change the instruction, you MUST re-record the cassettes.
_FROZEN_ABSTAIN_SYSTEM = (
    "You are a contract field generator. Output ONLY the NDA field values as a "
    "single JSON object matching the schema. No markdown fences, no prose. "
    "CRITICAL: if the requested governing law is NOT one of the Jurisdiction enum "
    "values, set governingLaw to \"OTHER\" and put the requested law verbatim in "
    "governingLawRaw -- do NOT substitute a different jurisdiction."
)


class TestAbstainableCodegen:

    def _policies(self):
        return json.loads((_NDA_DIR / "abstain-policy.json").read_text())["policies"]

    def test_composed_system_byte_identical(self):
        """T1 (CRITICAL regression): _abstain_system() (governingLaw, the default) composed
        from the generated policy is byte-identical to the frozen string -> Gauntlet replay
        key stable for the existing governing-law cassettes."""
        from contract_drafting import demo_mars_beat as demo
        demo.clear_cache()
        assert demo._abstain_system() == _FROZEN_ABSTAIN_SYSTEM
        assert demo._abstain_system("governingLaw") == _FROZEN_ABSTAIN_SYSTEM
        assert len(demo._abstain_system()) == 357

    def test_policy_structural_governing_law(self):
        """T2: the governingLaw policy is internally consistent and matches the schema enum."""
        gl = self._policies()["governingLaw"]
        assert gl["sentinel"] == "OTHER"
        assert gl["rawField"] == "governingLawRaw"
        assert gl["enum"] == "Jurisdiction"
        assert gl["instruction"].strip()
        schema = json.loads((_NDA_DIR / "schema.json").read_text())
        enum_vals = None
        for defn in schema.get("definitions", {}).values():
            if isinstance(defn, dict) and "enum" in defn and "Washington" in defn["enum"]:
                enum_vals = set(defn["enum"])
                break
        assert enum_vals is not None
        assert set(gl["representable"]) == enum_vals - {"OTHER"}
        assert len(gl["representable"]) == 64
        assert "OTHER" not in gl["representable"]

    def test_policy_structural_cross_field(self):
        """T2 cross-field: the new typed fields each get a consistent policy entry whose
        representable set is its enum MINUS its sentinel (the cross-field replication core)."""
        policies = self._policies()
        expected = {
            "disclosingEntityType": ("OTHER_ENTITY", "disclosingEntityTypeRaw", "EntityType", 11),
            "receivingEntityType": ("OTHER_ENTITY", "receivingEntityTypeRaw", "EntityType", 11),
            "disputeForum": ("OTHER_FORUM", "disputeForumRaw", "DisputeForum", 9),
        }
        for field, (sentinel, raw, enum, n_rep) in expected.items():
            p = policies[field]
            assert p["sentinel"] == sentinel
            assert p["rawField"] == raw
            assert p["enum"] == enum
            assert len(p["representable"]) == n_rep
            assert sentinel not in p["representable"]
            assert p["instruction"].strip()

    def test_instruction_matches_cto_decorator(self):
        """T2b: EVERY policy instruction is the verbatim @Abstainable 2nd arg in the
        .cto -- single source, no Python-side massaging, for all abstainable fields."""
        cto = (_NDA_DIR / "model" / "model.cto").read_text().replace('\\"', '"')
        for field, p in self._policies().items():
            assert p["instruction"] in cto, f"{field} instruction not verbatim in .cto"

    def test_cache_key_schema_strips_decorators(self):
        """Cache-key invariance (Codex): the schema used for the cassette key has
        NO $decorators, so adding @Abstainable to the .cto (which surfaces as
        schema.json $decorators) provably cannot perturb the replay key."""
        from contract_drafting import demo_mars_beat as demo
        schema = demo._field_schema(with_abstain=True)
        blob = json.dumps(schema)
        assert "$decorators" not in blob
        assert "Abstainable" not in blob

    def test_fail_closed_on_missing_policy(self, monkeypatch, tmp_path):
        """T4: a missing policy raises (never silently falls back to a hardcoded prompt)."""
        from contract_drafting import demo_mars_beat as demo
        from contract_drafting import schema_validator as sv
        demo.clear_cache()
        # Point the template base at an empty dir -> no abstain-policy.json -> {} ->
        # _abstain_policy("governingLaw") fails closed on the missing field.
        monkeypatch.setattr(sv, "_TEMPLATES_BASE", tmp_path)
        with pytest.raises(RuntimeError, match="no @Abstainable policy"):
            demo._abstain_system()
        demo.clear_cache()  # restore for other tests

    def test_import_safety_disclaimer_decoupled(self):
        """T5: the production draft path imports demo_mars_beat for LEGAL_DISCLAIMER;
        accessing it must NOT trigger the (eval-only) abstain-policy read."""
        from contract_drafting import demo_mars_beat as demo
        from contract_drafting import schema_validator as sv
        demo.clear_cache()
        _ = demo.LEGAL_DISCLAIMER  # production import path touches this, not the policy
        assert sv._policy_cache == {}, "disclaimer access must not load the abstain policy"

    def test_cli_prints_disclaimer_on_pass(self, monkeypatch, capsys, tmp_path):
        """The default (non-JSON) draft CLI must print the not-legal-advice footer on a
        successful draft -- the Ethics section's 'the UI states this' mitigation (Codex P2)."""
        import sys as _sys
        from contract_drafting import main as cdmain
        from contract_drafting.demo_mars_beat import LEGAL_DISCLAIMER

        monkeypatch.setattr(_sys, "argv", [
            "contract-drafting", "--mode", "draft",
            "--party-a", "TestCo", "--party-b", "AcmeCorp",
            "--effective-date", "2026-01-15",
            "--output-path", str(tmp_path / "cli.docx"),
            "--db-path", str(tmp_path / "cli.db"),
        ])
        with pytest.raises(SystemExit) as exc:
            cdmain.main()
        assert exc.value.code == 0
        assert LEGAL_DISCLAIMER in capsys.readouterr().out

    def test_cli_no_disclaimer_on_escalated(self, monkeypatch, capsys, tmp_path):
        """The PASS-specific footer must NOT print under a non-PASS gate (Codex P2):
        a 120-month term ESCALATES, so the disclaimer (which says 'PASS certifies...')
        would contradict the gate."""
        import sys as _sys
        from contract_drafting import main as cdmain
        from contract_drafting.demo_mars_beat import LEGAL_DISCLAIMER

        monkeypatch.setattr(_sys, "argv", [
            "contract-drafting", "--mode", "draft",
            "--party-a", "TestCo", "--party-b", "AcmeCorp",
            "--effective-date", "2026-01-15", "--term-months", "120",
            "--output-path", str(tmp_path / "cli.docx"),
            "--db-path", str(tmp_path / "cli.db"),
        ])
        with pytest.raises(SystemExit):
            cdmain.main()
        out = capsys.readouterr().out
        assert "ESCALATED" in out
        assert LEGAL_DISCLAIMER not in out

    def test_cli_typed_abstention_prints_escalated_not_bare_error(
            self, monkeypatch, capsys, tmp_path):
        """Codex P2: a typed-abstention result keeps an 'error' key (back-compat), but
        the CLI must surface the ESCALATED status, the abstained field with its raw
        ask, the audit_id, and the audit-lookup hint -- never the bare generic
        'Error: ...' that hides the escalation record."""
        import re as _re
        import sys as _sys
        from contract_drafting import main as cdmain

        monkeypatch.setattr(_sys, "argv", [
            "contract-drafting", "--mode", "draft",
            "--party-a", "TestCo", "--party-b", "AcmeCorp",
            "--effective-date", "2026-01-15",
            "--jurisdiction", "OTHER",
            "--output-path", str(tmp_path / "cli.docx"),
            "--db-path", str(tmp_path / "cli.db"),
        ])
        with pytest.raises(SystemExit) as exc:
            cdmain.main()
        assert exc.value.code == 1  # non-success, same convention as BLOCKED drafts
        out = capsys.readouterr().out
        assert "Gate: ESCALATED" in out
        assert "governingLaw" in out            # the abstained field is named
        assert _re.search(r"Audit ID: \d+", out)  # the audit row is surfaced ...
        assert _re.search(r"audit --doc \d+", out)  # ... with the lookup hint
        assert "Error:" not in out              # generic error branch did NOT fire

    def test_cli_json_redacts_rendered_text(self, monkeypatch, capsys, tmp_path):
        """C2: --json must NOT leak the full contract body (rendered_text) into the
        CLI JSON surface; render_sha256 stays as the verifiable link to the bytes."""
        import sys as _sys
        from contract_drafting import main as cdmain

        monkeypatch.setattr(_sys, "argv", [
            "contract-drafting", "--mode", "draft", "--json",
            "--party-a", "TestCo", "--party-b", "AcmeCorp",
            "--effective-date", "2026-01-15",
            "--output-path", str(tmp_path / "cli.docx"),
            "--db-path", str(tmp_path / "cli.db"),
        ])
        cdmain.main()  # --json path prints and returns (no sys.exit)
        out = capsys.readouterr().out
        payload = json.loads(out)
        assert payload["gate_result"] == "PASS"
        assert "render_sha256" in payload
        assert "rendered_text" not in payload
        assert "rendered_text" not in out

    def test_demo_format_uses_gate_neutral_disclaimer(self):
        """Codex P2: the Mars-beat demo only checks schema validity and both arms can be
        REJECTED, so its footer must be the gate-NEUTRAL line, never the PASS-specific
        'PASS certifies...' text."""
        from contract_drafting import demo_mars_beat as demo
        rejected = {"valid": False, "fields": {}, "errors": ["bad"]}
        out = demo._format({"free": rejected, "constrained": rejected})
        assert demo.NOT_LEGAL_ADVICE in out
        assert "PASS certifies" not in out

    @pytest.mark.skipif(not HAS_NODE, reason="node or @accordproject deps not installed")
    def test_regen_idempotent(self, tmp_path):
        """T3: re-running codegen produces byte-identical artifacts. Runs against a
        COPY of the committed template dir (the sibling tests' copytree pattern), so
        the working tree stays clean even if codegen output ever drifts; the copy's
        regenerated artifacts are compared byte-for-byte to the committed originals."""
        import subprocess
        dst = tmp_path / "nda-mutual"
        shutil.copytree(_NDA_DIR, dst)
        subprocess.run(
            ["node", "scripts/concerto-helper.js", str(dst)],
            cwd=_REPO_ROOT, check=True, capture_output=True,
        )
        for name in ("abstain-policy.json", "schema.json"):
            assert (dst / name).read_text() == (_NDA_DIR / name).read_text(), \
                f"npm run generate is not idempotent for {name}"

    @pytest.mark.skipif(not HAS_NODE, reason="node or @accordproject deps not installed")
    def test_regen_drops_one_fields_policy_but_keeps_others(self, tmp_path):
        """Codex P2 (multi-field): removing ONE field's @Abstainable -> that field's policy
        entry disappears on regen, the others remain (the sidecar tracks the .cto per field)."""
        import subprocess
        dst = tmp_path / "nda-mutual"
        shutil.copytree(_NDA_DIR, dst)
        cto_path = dst / "model" / "model.cto"
        # Drop ONLY governingLaw's field decorator (@Abstainable("OTHER", ...)).
        lines = [ln for ln in cto_path.read_text().splitlines(keepends=True)
                 if not ln.lstrip().startswith('@Abstainable("OTHER",')]
        cto_path.write_text("".join(lines))
        subprocess.run(["node", "scripts/concerto-helper.js", str(dst)],
                       cwd=_REPO_ROOT, check=True, capture_output=True)
        policies = json.loads((dst / "abstain-policy.json").read_text())["policies"]
        assert "governingLaw" not in policies
        assert "disclosingEntityType" in policies and "disputeForum" in policies

    @pytest.mark.skipif(not HAS_NODE, reason="node or @accordproject deps not installed")
    def test_regen_removes_stale_policy_when_all_decorators_absent(self, tmp_path):
        """Codex P2: when NO field is @Abstainable, regen DELETES a stale abstain-policy.json
        so the consumer fails closed rather than read outdated policy text."""
        import subprocess
        dst = tmp_path / "nda-mutual"
        shutil.copytree(_NDA_DIR, dst)
        cto_path = dst / "model" / "model.cto"
        # Drop ALL field-level @Abstainable(...) decorators (the bare enum markers
        # `@Abstainable` with no parens stay, but no FIELD is abstainable any more).
        lines = [ln for ln in cto_path.read_text().splitlines(keepends=True)
                 if not ln.lstrip().startswith('@Abstainable(')]
        cto_path.write_text("".join(lines))
        assert (dst / "abstain-policy.json").exists()
        subprocess.run(["node", "scripts/concerto-helper.js", str(dst)],
                       cwd=_REPO_ROOT, check=True, capture_output=True)
        assert not (dst / "abstain-policy.json").exists(), \
            "stale abstain-policy.json must be removed when no field is abstainable"


# ---------------------------------------------------------------------------
# TestRenderEscaping — P1.1: chevron HTML escaping disabled + markdown
# metacharacter policy. Party names / purposes must render VERBATIM into the
# canonical markdown and the signable docx ('Smith & Jones', never
# 'Smith &amp; Jones'), and free-text values must not be able to forge
# markdown block structure (headings/lists/HRs) through the pandoc path.
# ---------------------------------------------------------------------------

_AMPERSAND_PARTY = 'Smith & Jones "Alpha" LLC'


class TestRenderEscaping:

    def _ampersand_request(self) -> DraftRequest:
        return DraftRequest(
            disclosing_party=_AMPERSAND_PARTY,
            receiving_party="AcmeCorp",
            purpose="joint R&D",
            effective_date="2026-01-15",
        )

    def test_ampersand_and_quotes_render_literally(self):
        """End-to-end draft(): '&' and '\"' appear verbatim, never as HTML entities."""
        result = draft(self._ampersand_request())
        assert result.success, result.error
        assert _AMPERSAND_PARTY in result.text
        assert "joint R&D" in result.text
        assert "&amp;" not in result.text
        assert "&quot;" not in result.text
        assert "&lt;" not in result.text and "&gt;" not in result.text

    def test_draft_with_data_path_unescaped(self):
        """The raw-data render entry point (draft_with_data) is also unescaped."""
        from contract_drafting.cicero_bridge import draft_with_data
        data = {**_nda_valid_data(), "disclosingParty": _AMPERSAND_PARTY,
                "purpose": "joint R&D"}
        result = draft_with_data(data, template_name="nda-mutual")
        assert result.success, result.error
        assert _AMPERSAND_PARTY in result.text
        assert "joint R&D" in result.text
        assert "&amp;" not in result.text
        assert "&quot;" not in result.text

    def test_homebrew_docx_writes_unescaped_text(self, tmp_docx):
        """The pandoc-absent _homebrew_to_docx fallback (README-supported) writes the
        literal ampersand string into the signable docx, not '&amp;' (P1.1)."""
        from docx import Document
        result = draft(self._ampersand_request())
        assert result.success, result.error
        path = _homebrew_to_docx(result.text, tmp_docx)
        full_text = "\n".join(p.text for p in Document(path).paragraphs)
        assert _AMPERSAND_PARTY in full_text
        assert "joint R&D" in full_text
        assert "&amp;" not in full_text
        assert "&quot;" not in full_text

    @pytest.mark.skipif(not HAS_PANDOC, reason="pandoc not installed")
    def test_pandoc_docx_writes_unescaped_text(self, tmp_docx):
        """The pandoc path also produces the literal '&' in the docx."""
        from docx import Document
        result = draft(self._ampersand_request())
        path = markdown_to_docx(result.text, tmp_docx)
        full_text = "\n".join(p.text for p in Document(path).paragraphs)
        assert 'Smith & Jones' in full_text
        assert "&amp;" not in full_text

    def test_shared_clause_partials_unescaped(self):
        """Variables INSIDE shared-clause partials are unescaped too (the no-escape
        preprocessing covers partials, not just the top-level grammar)."""
        from contract_drafting.cicero_bridge import draft_with_data
        data = {**_jv_test_data(), "partyAName": "R&D Holdings Pte. Ltd."}
        result = draft_with_data(data, template_name="joint-venture")
        assert result.success, result.error
        assert "R&D Holdings Pte. Ltd." in result.text
        assert "&amp;" not in result.text

    def test_newline_injection_cannot_forge_heading(self):
        """A party name with embedded newlines cannot start a markdown block: the
        newlines are collapsed, so '# INJECTED' never lands at a line start."""
        req = DraftRequest(
            disclosing_party="Evil\n# INJECTED HEADING\n- fake list\nCo",
            receiving_party="AcmeCorp", effective_date="2026-01-15",
        )
        result = draft(req)
        assert result.success, result.error
        assert "\n# INJECTED" not in result.text
        assert "\n- fake list" not in result.text
        # The content survives inline (collapsed), it just cannot be structural.
        assert "Evil # INJECTED HEADING - fake list Co" in result.text

    def test_sanitize_md_scalar_policy(self):
        """Unit spec for the markdown-metacharacter policy: newline collapse +
        leading block-marker escape; plain values are byte-identical no-ops."""
        from contract_drafting.cicero_bridge import _sanitize_md_scalar
        # No-ops: ordinary names, phone numbers, decimals (byte-stability of
        # existing data_hash / render_sha256 values depends on this).
        for value in ("Acme Corporation", "+65 0000 0000", "3.5x growth",
                      "Smith & Jones \"Alpha\" LLC", "a-b_c*d"):
            assert _sanitize_md_scalar(value) == value
        # Newline runs collapse to a single space.
        assert _sanitize_md_scalar("A\nB\r\nC") == "A B C"
        # Leading block markers that WOULD open a block get a backslash.
        assert _sanitize_md_scalar("# Heading Co") == "\\# Heading Co"
        assert _sanitize_md_scalar("- List Co") == "\\- List Co"
        assert _sanitize_md_scalar("* List Co") == "\\* List Co"
        assert _sanitize_md_scalar("+ List Co") == "\\+ List Co"
        assert _sanitize_md_scalar("> Quote Co") == "\\> Quote Co"
        assert _sanitize_md_scalar("1. Ordered Co") == "1\\. Ordered Co"
        assert _sanitize_md_scalar("---") == "\\---"
        # Markers that would NOT open a block are left alone.
        assert _sanitize_md_scalar("#hashtag") == "#hashtag"
        assert _sanitize_md_scalar("1.5 million") == "1.5 million"

    def test_disable_html_escaping_tag_rewrite(self):
        """Unit spec for the no-escape preprocessing: plain {{var}} becomes {{&var}};
        section/close/inverted/partial/comment/set-delimiter and already-unescaped
        tags are untouched."""
        from contract_drafting.cicero_bridge import _disable_html_escaping
        assert _disable_html_escaping("{{name}}") == "{{&name}}"
        assert _disable_html_escaping("{{ name }}") == "{{& name }}"
        for untouched in ("{{#cond}}x{{/cond}}", "{{^cond}}x{{/cond}}",
                          "{{> partial}}", "{{&raw}}", "{{{raw}}}",
                          "{{! comment}}", "{{=| |=}}"):
            assert _disable_html_escaping(untouched) == untouched

    def test_escaped_draft_hash_consistency(self):
        """data_hash covers the SANITIZED data (what the renderer consumed), so
        audit_id -> rendered bytes stays 1:1; plain inputs hash unchanged."""
        from contract_drafting.cicero_bridge import _sanitize_md_fields
        r1 = draft(DraftRequest(disclosing_party="TestCo", receiving_party="AcmeCorp",
                                effective_date="2026-01-15"))
        d = _build_data(DraftRequest(disclosing_party="TestCo", receiving_party="AcmeCorp",
                                     effective_date="2026-01-15"))
        assert _sanitize_md_fields(d) == d  # plain input: sanitize is a no-op
        assert r1.data_hash == _compute_data_hash(d)


# ---------------------------------------------------------------------------
# TestRefineForbiddenSyntax — P1.2b: LLM-refined clause text must not smuggle
# {{> partial}}, {{&var}}, {{{var}}}, or {{=delim=}} through validation (these
# evade the variable extractor and render as silent injection / dropped text).
# ---------------------------------------------------------------------------

class TestRefineForbiddenSyntax:

    def _mock_llm_response(self, monkeypatch, response: str) -> None:
        monkeypatch.setattr(refine_mod, "_call_refine_llm",
                            lambda system_prompt, user_prompt: response)

    @pytest.mark.parametrize("payload", [
        # Partial injection: renders another clause (or silently drops).
        "**Applicable Law:** {{governingLaw}} law applies.\n\n{{> signature-block}}",
        # Unescaped-variable tag: bypasses the allowed-vars check entirely.
        "**Applicable Law:** {{&governingLaw}} law applies.",
        # Triple mustache: same unescaped channel.
        "**Applicable Law:** {{{governingLaw}}} law applies.",
        # Delimiter change: every LATER tag would evade all the regex layers.
        "{{=| |=}} **Applicable Law:** |governingLaw| law applies.",
    ])
    def test_forbidden_constructs_blocked(self, monkeypatch, payload):
        self._mock_llm_response(monkeypatch, payload)
        r = refine_clause("governing-law", "rewrite",
                          data={"governingLaw": "Singapore"},
                          template_name="joint-venture")
        assert r.status == "BLOCKED", f"payload passed validation: {payload!r}"
        assert any("forbidden Mustache syntax" in e for e in r.validation.errors)

    def test_forbidden_partial_with_unknown_name_blocked(self, monkeypatch):
        """A partial referencing a NON-existent clause (the silent-drop case:
        chevron renders it as empty text) is rejected the same way."""
        self._mock_llm_response(
            monkeypatch,
            "**Applicable Law:** {{governingLaw}} law applies. {{> no-such-clause}}",
        )
        r = refine_clause("governing-law", "rewrite",
                          data={"governingLaw": "Singapore"},
                          template_name="joint-venture")
        assert r.status == "BLOCKED"
        assert any("forbidden Mustache syntax" in e for e in r.validation.errors)

    def test_validate_refined_clause_unit(self):
        """Direct unit check on validate_refined_clause (no LLM plumbing)."""
        from contract_drafting.refine import validate_refined_clause
        bad = validate_refined_clause(
            "text {{> injected}} more {{&sneaky}} text",
            allowed_vars={"governingLaw"}, original_text="text",
        )
        assert not bad.ok
        assert any("forbidden Mustache syntax" in e for e in bad.errors)
        good = validate_refined_clause(
            "**Law:** {{governingLaw}} applies.",
            allowed_vars={"governingLaw"}, original_text="text",
        )
        assert good.ok, good.errors

    def test_canonical_clauses_would_pass_forbidden_check(self):
        """No canonical shared clause uses a forbidden construct, so the new layer
        cannot block a legitimate refinement round-trip."""
        from contract_drafting.refine import _FORBIDDEN_TAG_RE, SHARED_CLAUSES_DIR
        for p in sorted(SHARED_CLAUSES_DIR.glob("*.md")):
            assert not _FORBIDDEN_TAG_RE.search(p.read_text(encoding="utf-8")), (
                f"canonical clause {p.name} contains a forbidden construct; "
                f"the P1.2b check would need a carve-out"
            )


# ---------------------------------------------------------------------------
# TestBlockedDraftNoArtifact — P1.2a: a playbook-BLOCKED draft on the legacy
# --engine llm path must NOT leave a signable .docx on disk (previously it was
# assembled 'for review' and then disowned by the audit row).
# ---------------------------------------------------------------------------

class TestBlockedDraftNoArtifact:

    def _mock_fields(self, monkeypatch):
        from contract_drafting import compliance_draft as cd
        monkeypatch.setattr(cd, "_call_llm_for_fields", lambda *a, **k: {
            "purpose": "x", "term_months": 24, "notice_days": 30,
            "survival_years": 3, "governing_law": "New York",
        })

    def test_llm_blocked_leaves_no_docx(self, monkeypatch, tmp_db, tmp_docx):
        from contract_drafting import compliance_draft as cd
        self._mock_fields(monkeypatch)
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="New York",
            has_non_compete=True,  # playbook: prohibited -> BLOCKED
            output_path=tmp_docx,
        )
        result = draft_contract(req, engine="llm", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert result["output_path"] is None
        assert not Path(tmp_docx).exists(), \
            "BLOCKED draft left a signable artifact on disk (fail-closed violation)"
        assert any(v["clause_type"] == "Non-Compete" for v in result["violations"])
        # Audit row agrees: BLOCKED, no output_path, violations recorded.
        rows = cd.get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert rows and rows[0]["gate_result"] == "BLOCKED"
        assert rows[0]["output_path"] is None
        assert "Non-Compete" in (rows[0]["violations"] or "")

    def test_llm_blocked_default_output_path_no_orphan(self, monkeypatch, tmp_db, tmp_path):
        """Same fail-closed property when the caller supplies no output_path (the
        default data/drafts/... path must not be written either)."""
        import glob
        self._mock_fields(monkeypatch)
        before = set(glob.glob("data/drafts/*.docx"))
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", governing_law="New York",
            has_non_solicitation=True,  # playbook: prohibited -> BLOCKED
        )
        result = draft_contract(req, engine="llm", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert result["output_path"] is None
        assert set(glob.glob("data/drafts/*.docx")) == before, \
            "BLOCKED draft wrote an orphan docx under data/drafts/"

    def test_cicero_blocked_leaves_no_docx(self, tmp_db, tmp_docx):
        """Parity regression guard: the cicero path already blocks before any
        render -- keep it that way."""
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-15", has_non_compete=True,
            output_path=tmp_docx,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        assert not Path(tmp_docx).exists()


# ---------------------------------------------------------------------------
# TestCiceroTemplateResolution — P1.2c: _draft_cicero resolves the template
# ONCE via the registry (aliases included), and every audit row carries the
# actual resolved template id (no hardcoded 'cicero/nda-mutual').
# ---------------------------------------------------------------------------

class TestCiceroTemplateResolution:

    @pytest.mark.parametrize("alias,expected", [
        ("jv", "joint-venture"),
        ("finder", "intermediary"),
        ("consultant", "consulting"),
        ("cooperation", "strategic-cooperation"),
    ])
    def test_registry_aliases_resolve(self, alias, expected):
        from contract_drafting.template_registry import get_registry
        info = get_registry().get_for_doc_type(alias)
        assert info is not None, f"alias {alias!r} not registered"
        assert info.name == expected
        assert info.path.is_dir()

    def test_alias_draft_uses_registry_directory(self, tmp_db):
        """doc_type='jv' flows through as 'joint-venture' end-to-end: the schema
        pre-gate validates against the JV schema (NDA-shaped request -> BLOCKED
        with JV field errors), and the audit row names cicero/joint-venture --
        never the raw alias 'cicero/jv' or a second divergent resolution."""
        from contract_drafting import compliance_draft as cd
        req = DraftRequest(
            doc_type="jv", disclosing_party="A", receiving_party="B",
            effective_date="2026-01-01",
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        # Blocked by the JOINT-VENTURE schema (proof the right dir was resolved),
        # not by a missing 'jv' template directory.
        assert "Template directory not found" not in result.get("error", "")
        rows = cd.get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert rows and rows[0]["template_id"] == "cicero/joint-venture"

    def test_playbook_blocked_audit_row_names_actual_template(self, tmp_db):
        """The playbook-BLOCKED audit row uses the registry-resolved template id
        (for an NDA that is cicero/nda-mutual -- previously hardcoded; this pins
        the f-string form so the JV/consulting rows cannot regress silently)."""
        from contract_drafting import compliance_draft as cd
        req = DraftRequest(
            disclosing_party="TestCo", receiving_party="AcmeCorp",
            effective_date="2026-01-01", has_non_compete=True,
        )
        result = draft_contract(req, engine="cicero", db_path=tmp_db)
        assert result["gate_result"] == "BLOCKED"
        rows = cd.get_audit_log(tmp_db, doc_id=result["audit_id"])
        assert rows and rows[0]["template_id"] == "cicero/nda-mutual"
