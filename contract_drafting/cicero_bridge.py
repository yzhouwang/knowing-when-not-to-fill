"""
cicero_bridge.py — CiceroMark template engine bridge.

Renders CiceroMark templates (Mustache-syntax markdown with {{variables}} and
{{#if condition}} blocks) deterministically from structured data. No LLM in the
draft loop — same inputs always produce the same contract.

Uses chevron (Mustache renderer) for template processing and python-docx for
optional markdown-to-docx conversion.
"""
from __future__ import annotations

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

import chevron

if TYPE_CHECKING:
    from contract_drafting.compliance_draft import DraftRequest

log = logging.getLogger(__name__)

# Default template base directory (relative to project root)
_TEMPLATES_BASE = Path(__file__).resolve().parent.parent / "data" / "templates" / "cicero"

# Mustache partial anchor syntax: {{> clause-name}} or {{>clause-name}}
_ANCHOR_RE = re.compile(r"\{\{>\s*([a-z][a-z0-9-]*)\s*\}\}")

# P1.1: chevron 2.x HTML-escapes every plain {{variable}} tag ('&' -> '&amp;',
# '"' -> '&quot;', '<'/'>' -> '&lt;'/'&gt;'). That is wrong for contracts: the
# canonical rendered markdown (hashed into the audit row as render_sha256) and
# the pandoc-absent _homebrew_to_docx fallback would both carry HTML entities
# literally ('Smith & Jones' -> 'Smith &amp; Jones' in a signable docx).
# CONSTRAINT: the installed chevron.render() exposes NO escape kwarg (its
# signature is template/data/partials only), and monkeypatching
# chevron.renderer._html_escape would be process-global. So the fix is template
# preprocessing at render time: rewrite every plain {{var}} interpolation tag to
# the Mustache no-escape form {{&var}}. Section/inverted/close/partial/comment/
# set-delimiter tags ({{#..}} {{^..}} {{/..}} {{>..}} {{!..}} {{=..=}}) and
# already-unescaped tags ({{&..}} {{{..}}}) are left untouched (the lookbehind
# keeps the scan from re-matching inside a triple mustache). Code-side, so
# EVERY template and shared-clause partial is covered without editing templates.
_PLAIN_VARIABLE_TAG_RE = re.compile(r"(?<!\{)\{\{(?![#/^>&!{=])")

# P1.1 markdown-metacharacter policy: with HTML escaping off, a free-text value
# (party name, purpose) flows verbatim into the rendered markdown that pandoc
# interprets. Block-level injection (a heading, list, blockquote, or HR forged
# from a party name) requires the value to start a LINE, which it can only do
# via an embedded newline (template variables sit mid-sentence) or -- defensively
# -- by beginning with a structural marker while rendered at a line start. So:
# collapse newline/CR runs to a single space, and backslash-escape a leading
# block marker (#, >, +, *, -, or an ordered-list '1.'/'1)'). Pandoc unescapes
# the backslash; the homebrew fallback only treats '#'/'---' at a full-line
# start, which the newline collapse already prevents.
# RESIDUAL (accepted, documented): inline metacharacters (*emphasis*,
# _underscore_, [link](url), `code`) are NOT escaped -- they can restyle text
# within the surrounding paragraph but cannot create block structure, and
# escaping them would put literal backslashes into the canonical rendered
# markdown and the homebrew docx fallback (which performs no unescaping).
# Markers are escaped only when they would ACTUALLY open a block (heading/list
# markers need a following space in pandoc markdown; '>' does not; an HR is a
# bare run of -/*/_), so ordinary values like '+65 0000 0000' or '3.5x' stay
# byte-identical.
_MD_NEWLINE_RUN_RE = re.compile(r"[\r\n]+")
_MD_LEADING_BLOCK_MARKER_RE = re.compile(
    r"^(\s*)(#{1,6}(?=\s|$)|[+*-](?=\s|$)|[-*_]{3,}(?=\s|$)|>)"
)
_MD_LEADING_ORDERED_LIST_RE = re.compile(r"^(\s*\d{1,9})([.)])(?=\s|$)")

# Module-level partials cache: shared_dir -> (latest_mtime, partials_dict)
_partials_cache: dict[Path, tuple[float, dict[str, str]]] = {}


@dataclass
class CiceroDraftResult:
    """Result of a CiceroMark template draft."""
    text: str                # Generated contract markdown
    template_name: str       # e.g. "nda-mutual"
    template_version: str    # From package.json
    data_hash: str           # SHA256 of input data (audit reproducibility)
    success: bool
    error: str | None = None
    # Non-fatal diagnostics surfaced to the caller (e.g. a captured-but-not-rendered
    # dispute forum that a human must resolve). The draft still succeeds; warnings are
    # NOT part of data_hash, so they never perturb audit reproducibility.
    warnings: list[str] = field(default_factory=list)


def _disable_html_escaping(template_text: str) -> str:
    """Rewrite plain {{var}} tags to the no-escape form {{&var}} (P1.1).

    See _PLAIN_VARIABLE_TAG_RE for why this exists (installed chevron has no
    escape switch) and what it deliberately leaves alone.
    """
    return _PLAIN_VARIABLE_TAG_RE.sub("{{&", template_text)


def _render_unescaped(template_text: str, data: dict, partials: dict[str, str]) -> str:
    """chevron.render with HTML escaping disabled for the template AND every partial.

    The single render chokepoint for both draft entry points, so no call site can
    reintroduce entity-escaped output (P1.1).
    """
    return chevron.render(
        _disable_html_escaping(template_text),
        data,
        partials_dict={name: _disable_html_escaping(text) for name, text in partials.items()},
    )


def _sanitize_md_scalar(value: str) -> str:
    """Neutralize markdown BLOCK-structure injection in one free-text value (P1.1).

    No-op for ordinary names/purposes (no newlines, no leading block marker), so
    existing data_hash / render_sha256 values are unchanged for all committed
    eval and demo inputs. See the _MD_* regex comment block for the policy and
    the documented inline-metacharacter residual.
    """
    value = _MD_NEWLINE_RUN_RE.sub(" ", value)
    value = _MD_LEADING_BLOCK_MARKER_RE.sub(r"\1\\\2", value)
    return _MD_LEADING_ORDERED_LIST_RE.sub(r"\1\\\2", value)


def _sanitize_md_fields(data: dict) -> dict:
    """Apply _sanitize_md_scalar to every top-level string value in render data."""
    return {
        key: _sanitize_md_scalar(value) if isinstance(value, str) else value
        for key, value in data.items()
    }


def _resolve_template_dir(template_name: str) -> Path:
    """Resolve template name to directory path."""
    template_dir = _TEMPLATES_BASE / template_name
    if not template_dir.is_dir():
        raise FileNotFoundError(f"Template directory not found: {template_dir}")
    return template_dir


def _read_template(template_dir: Path) -> str:
    """Read the CiceroMark grammar template. Pure — no composition."""
    grammar_path = template_dir / "text" / "grammar.tem.md"
    if not grammar_path.exists():
        raise FileNotFoundError(f"Template file not found: {grammar_path}")
    return grammar_path.read_text(encoding="utf-8")


def _scan_anchors(template_text: str) -> set[str]:
    """Extract the set of {{> clause-name}} partial anchor names in a template.

    Matches Mustache partial syntax. Deduplicates repeated anchors.
    """
    return set(_ANCHOR_RE.findall(template_text))


def _load_partials(
    shared_dir: Path,
    required: set[str],
    *,
    overlay_map: dict[str, str] | None = None,
    clause_overrides: dict[str, str] | None = None,
) -> dict[str, str]:
    """Load the required shared-clause partials from disk, with mtime-invalidated cache.

    Resolution order (first match wins) for each required partial:
    1. clause_overrides[name] — ephemeral in-memory override (refinement preview)
    2. overlay_map[name] → _overlays/{name}/{overlay_id}.md — library-promoted variant
    3. shared-clauses/{name}.md — canonical default

    Raises FileNotFoundError if:
    - shared_dir does not exist (and required is non-empty)
    - any required partial is missing from the canonical directory AND not overridden
    - an overlay_map entry points to a non-existent overlay file

    The cache is keyed by shared_dir and invalidates when any *.md in the dir
    changes mtime (catches adds/edits/deletes during development). Overlays and
    ephemeral overrides are NOT cached (they're per-call).
    """
    if not required:
        return {}
    if not shared_dir.is_dir():
        raise FileNotFoundError(f"Shared clauses directory not found: {shared_dir}")

    overlay_map = overlay_map or {}
    clause_overrides = clause_overrides or {}

    # Load canonical partials from the cache first so we have fallback.
    md_files = list(shared_dir.glob("*.md"))
    latest_mtime = max((p.stat().st_mtime for p in md_files), default=0.0)
    cached = _partials_cache.get(shared_dir)
    if cached is not None and cached[0] == latest_mtime:
        all_canonical = cached[1]
    else:
        all_canonical = {p.stem: p.read_text(encoding="utf-8") for p in md_files}
        _partials_cache[shared_dir] = (latest_mtime, all_canonical)

    result: dict[str, str] = {}
    missing: list[str] = []
    overlays_dir = shared_dir / "_overlays"

    for name in required:
        # 1. Ephemeral override wins
        if name in clause_overrides:
            result[name] = clause_overrides[name]
            continue
        # 2. Overlay lookup
        if name in overlay_map:
            overlay_id = overlay_map[name]
            overlay_path = overlays_dir / name / f"{overlay_id}.md"
            if not overlay_path.is_file():
                raise FileNotFoundError(
                    f"Overlay not found: {name}/{overlay_id} "
                    f"(expected at {overlay_path})"
                )
            result[name] = overlay_path.read_text(encoding="utf-8")
            continue
        # 3. Canonical fallback
        if name in all_canonical:
            result[name] = all_canonical[name]
            continue
        missing.append(name)

    if missing:
        raise FileNotFoundError(
            f"Template references missing shared clauses: {sorted(missing)}. "
            f"Available in {shared_dir}: {sorted(all_canonical.keys())}"
        )
    return result


def _clear_partials_cache() -> None:
    """Clear the module-level partials cache. Used by tests and dev workflows."""
    _partials_cache.clear()


def _normalize_array_fields(data: dict) -> dict:
    """Auto-derive `has<Field>` and `<field>Joined` for every list[str] in data.

    Mustache section iteration on a list of bare strings always produces leading
    or trailing separators (Phase 0 spike confirmed). The clean pattern is a
    conditional section + pre-joined string:

        {{#hasLocalComplianceLaws}}, {{localComplianceLawsJoined}}{{/hasLocalComplianceLaws}}

    For any field `foo: list[str]`, this helper adds:
      - `hasFoo: bool(value)`     — drives the conditional opener
      - `fooJoined: ", ".join(value)` — the rendered string

    Original list field is preserved (templates may iterate manually if needed).
    Empty lists yield `hasFoo=False` so the section is suppressed entirely.

    Non-list fields and lists of non-strings are left untouched. Idempotent:
    re-running on already-normalized data is a no-op.

    Merge order: derived fields are FALLBACKS — if the caller already supplied
    `hasLocalComplianceLaws=False` to explicitly suppress the conditional even
    though the list is non-empty, that explicit value wins over the auto-derived
    one. `{**derived, **data}` puts data last so user intent is preserved.
    """
    derived: dict = {}
    for key, value in data.items():
        if not isinstance(value, list):
            continue
        # Only normalize lists of strings; skip lists of dicts / mixed.
        if value and not all(isinstance(item, str) for item in value):
            continue
        # Skip already-derived fields (defensive: don't double-normalize)
        if key.startswith("has") or key.endswith("Joined"):
            continue
        cap = key[:1].upper() + key[1:]
        derived[f"has{cap}"] = bool(value)
        derived[f"{key}Joined"] = ", ".join(value)
    return {**derived, **data}


def _read_composition_variants(template_dir: Path) -> dict:
    """Read variant flags from composition.json (merged into template data)."""
    comp_path = template_dir / "composition.json"
    if not comp_path.exists():
        return {}
    comp = json.loads(comp_path.read_text(encoding="utf-8"))
    return comp.get("variants", {})


def _read_template_version(template_dir: Path) -> str:
    """Read the template version from package.json."""
    pkg_path = template_dir / "package.json"
    if not pkg_path.exists():
        return "0.0.0"
    try:
        pkg = json.loads(pkg_path.read_text(encoding="utf-8"))
        return pkg.get("version", "0.0.0")
    except (json.JSONDecodeError, OSError):
        return "0.0.0"


def _to_display_governing_law(value: str, template_name: str) -> str:
    """Map a governingLaw value to its display name for rendering.

    Accepts an identifier ("New_York") or an already-display name ("Republic of
    Singapore"). FAIL-CLOSED: a genuinely unknown identifier raises, so we never
    render an unrecognized jurisdiction into a contract. For a template with no
    Jurisdiction map, jurisdiction_map is identity, so the value is unchanged.
    """
    from contract_drafting import jurisdiction_map

    if value == "OTHER":
        # PR2 abstain sentinel: NEVER render OTHER as a governing law. Fail closed here
        # (the render chokepoint) so EVERY path -- draft(), draft_with_data(), and any
        # future caller -- blocks; the draft entry points additionally pre-check OTHER
        # for a cleaner BLOCKED result with the raw ask surfaced.
        raise ValueError(
            "governingLaw=OTHER is the abstain sentinel and must not be rendered; a human "
            "must supply a supported governing law (see governingLawRaw)."
        )

    try:
        return jurisdiction_map.to_display(value, template_name=template_name)
    except ValueError:
        # Not a known identifier; accept only if it's a known display name.
        if jurisdiction_map.to_identifier(value, template_name=template_name) == value:
            raise  # neither a valid identifier nor a known display name
        return value


def _to_display_entity_type(
    value: str,
    template_name: str,
    field_name: str = "disclosingEntityType/receivingEntityType",
) -> str:
    """Map an entityType value to its display name for rendering. Accepts an identifier
    ("limited_liability_company") or an already-display name ("limited liability company").
    FAIL-CLOSED: OTHER_ENTITY is the abstain sentinel and must never render; an unknown
    value raises. For a template with no EntityType map, returns the value verbatim.

    field_name parametrizes the error message so callers that know WHICH field carried
    the sentinel (draft_with_data's raw-data path) name it precisely (M5)."""
    from contract_drafting import jurisdiction_map
    if value == "OTHER_ENTITY":
        raise ValueError(
            f"{field_name}=OTHER_ENTITY is the abstain sentinel "
            f"and must not be rendered; a human must supply a supported entity form "
            f"(see the *EntityTypeRaw field)."
        )
    return jurisdiction_map.to_display_enum("EntityType", value, template_name)


def _forum_capture_warnings(forum, raw=None) -> list[str]:
    """Non-fatal warnings for a captured-but-not-rendered dispute forum (Codex P2 / T6).

    disputeForum is a captured/audit field, never auto-inserted into the rendered clause
    (rendering a speculatively-filled forum would silently change a material legal term --
    court vs arbitration -- without user intent). But a captured forum must NOT be silently
    dropped either: this returns the human-readable warning(s) the caller surfaces. Empty
    list when no forum was supplied, so normal drafts stay warning-free. Shared by BOTH
    render entry points (draft / draft_with_data) so the wording cannot drift."""
    warnings: list[str] = []
    if isinstance(forum, str) and forum:
        if forum == "OTHER_FORUM":
            _detail = f": '{raw}'" if isinstance(raw, str) and raw else ""
            warnings.append(
                f"disputeForum=OTHER_FORUM (abstain sentinel{_detail}) was captured but is NOT "
                f"rendered into the clause; the draft retains the default court venue. A human "
                f"must resolve the dispute-resolution forum before execution."
            )
        else:
            warnings.append(
                f"A dispute-resolution forum ({forum}) was requested but disputeForum is a "
                f"captured/audit field, NOT auto-inserted into the clause (auto-rendering would "
                f"silently flip a material term, court vs arbitration); the draft retains the "
                f"default court venue. Resolve the forum manually if arbitration is intended."
            )
    return warnings


def _build_data(request: "DraftRequest", template_name: str = "nda-mutual") -> dict:
    """Map DraftRequest (snake_case) to template variables (camelCase).

    template_name selects which Jurisdiction map to apply. Only templates whose
    model.cto declares the native Jurisdiction enum have a map; for all others
    jurisdiction_map is an identity no-op, so governingLaw renders verbatim.
    """
    effective_date = request.effective_date or datetime.now().strftime("%Y-%m-%d")
    governing_law = _to_display_governing_law(request.governing_law, template_name)

    return {
        "effectiveDate": effective_date,
        "disclosingParty": request.disclosing_party,
        "receivingParty": request.receiving_party,
        "disclosingEntityType": _to_display_entity_type(request.disclosing_entity_type, template_name),
        "receivingEntityType": _to_display_entity_type(request.receiving_entity_type, template_name),
        "purpose": request.purpose,
        "termMonths": request.term_months,
        "noticeDays": request.notice_days,
        "survivalYears": request.survival_years,
        "governingLaw": governing_law,
        "mutual": request.mutual,
        "hasNonCompete": request.has_non_compete,
        "hasNonSolicitation": request.has_non_solicitation,
        "hasResidualsClause": request.has_residuals_clause,
    }


def _compute_data_hash(
    data: dict,
    *,
    overlay_map: dict[str, str] | None = None,
    clause_overrides: dict[str, str] | None = None,
) -> str:
    """Compute SHA256 hash of input data (and overlay/override identity) for audit.

    The hash must fold in overlay_map AND clause_overrides so that two drafts with
    identical input data but different partials produce distinct hashes. Otherwise
    audit_id → rendered bytes is no longer 1:1.
    """
    overlay_map = overlay_map or {}
    clause_overrides = clause_overrides or {}
    # Hash overrides by sha256 of their content so the hash is stable regardless
    # of whitespace; the content itself ends up in the rendered output anyway.
    overrides_hashed = {
        name: hashlib.sha256(text.encode("utf-8")).hexdigest()
        for name, text in clause_overrides.items()
    }
    canonical = json.dumps(
        {
            "data": data,
            "overlay_map": overlay_map,
            "clause_overrides_sha256": overrides_hashed,
        },
        sort_keys=True,
        default=str,
    )
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def draft(
    request: "DraftRequest",
    *,
    template_name: str = "nda-mutual",
    clause_overrides: dict[str, str] | None = None,
    overlay_map: dict[str, str] | None = None,
) -> CiceroDraftResult:
    """Render a CiceroMark template with data from a DraftRequest.

    Deterministic: same request always produces the same output.

    Composition: templates use {{> clause-name}} anchors to include shared
    clauses at exact positions. Partials are loaded from data/templates/shared-clauses/.
    If a template references a missing partial, draft() returns success=False
    with a clear error — no silent empty rendering.

    Args:
        clause_overrides: Ephemeral in-memory overrides {clause_name: raw_text}.
            Win over overlays and canonical. Used for refinement previews.
        overlay_map: Library-promoted overlays {clause_name: overlay_id} that
            resolve to _overlays/{clause_name}/{overlay_id}.md. Durable reference.
    """
    try:
        template_dir = _resolve_template_dir(template_name)
        template_text = _read_template(template_dir)
        version = _read_template_version(template_dir)

        # Scan for partial anchors and load them (raises if any missing)
        required_anchors = _scan_anchors(template_text)
        shared_dir = template_dir.parent.parent / "shared-clauses"
        partials = _load_partials(
            shared_dir, required_anchors,
            overlay_map=overlay_map,
            clause_overrides=clause_overrides,
        )

        # Build data and merge variant flags
        data = _build_data(request, template_name=template_name)
        variants = _read_composition_variants(template_dir)
        data.update(variants)
        data = _normalize_array_fields(data)
        # P1.1: sanitize BEFORE hashing so audit_id -> rendered bytes stays 1:1
        # (the hash covers exactly what the renderer consumed).
        data = _sanitize_md_fields(data)
        data_hash = _compute_data_hash(
            data,
            overlay_map=overlay_map,
            clause_overrides=clause_overrides,
        )

        rendered = _render_unescaped(template_text, data, partials)

        # T6: a dispute forum supplied on the request is captured-not-rendered (it is
        # deliberately NOT part of the render data, so data_hash/rendered bytes are
        # unchanged) -- surface the capture as a warning, never a silent drop.
        warnings = _forum_capture_warnings(
            getattr(request, "dispute_forum", ""),
            getattr(request, "dispute_forum_raw", ""),
        )

        return CiceroDraftResult(
            text=rendered,
            template_name=template_name,
            template_version=version,
            data_hash=data_hash,
            success=True,
            warnings=warnings,
        )
    except Exception as e:
        log.error(f"Cicero draft failed: {e}")
        return CiceroDraftResult(
            text="",
            template_name=template_name,
            template_version="0.0.0",
            data_hash="",
            success=False,
            error=str(e),
        )


def draft_with_data(
    data: dict,
    *,
    template_name: str,
    clause_overrides: dict[str, str] | None = None,
    overlay_map: dict[str, str] | None = None,
) -> CiceroDraftResult:
    """Render a CiceroMark template with a raw data dict.

    Used by agent flows that produce camelCase data directly (e.g. from plain-language
    parsing) without going through DraftRequest. Same composition + partial loading
    as draft().

    Args:
        clause_overrides: Ephemeral in-memory overrides {clause_name: raw_text}.
            Win over overlays and canonical. Used for refinement previews.
        overlay_map: Library-promoted overlays {clause_name: overlay_id} that
            resolve to _overlays/{clause_name}/{overlay_id}.md.
    """
    try:
        template_dir = _resolve_template_dir(template_name)
        template_text = _read_template(template_dir)
        version = _read_template_version(template_dir)

        required_anchors = _scan_anchors(template_text)
        shared_dir = template_dir.parent.parent / "shared-clauses"
        partials = _load_partials(
            shared_dir, required_anchors,
            overlay_map=overlay_map,
            clause_overrides=clause_overrides,
        )

        variants = _read_composition_variants(template_dir)
        full_data = _normalize_array_fields({**data, **variants})
        # Raw-data render path: map a governingLaw identifier ("New_York") to its
        # display name so a schema-valid identifier never renders verbatim into a
        # contract. Fail-closed on an unknown identifier; identity for templates
        # without a Jurisdiction map (the JV/consulting/etc. raw-data callers).
        if isinstance(full_data.get("governingLaw"), str):
            full_data["governingLaw"] = _to_display_governing_law(
                full_data["governingLaw"], template_name
            )
        # Display mapping for the RENDERED typed enum fields (entityType), fail-closed on the
        # abstain sentinel -- a sentinel must NEVER render as a real value. disputeForum is
        # deliberately NOT here: it is a captured/audit field, never auto-inserted into the
        # rendered clause (rendering a speculatively-filled forum would silently change a
        # material legal term -- court vs arbitration -- without user intent).
        #
        # Single-sourced through _to_display_entity_type (M5): it FAILS CLOSED on the
        # OTHER_ENTITY sentinel AND on an un-representable value (e.g. "GmbH", "KCAB") --
        # the raised ValueError propagates out of draft_with_data as a failed result, so
        # the render never silently emits an unsupported enum value. Identifiers and known
        # display names map cleanly; templates without the enum map are an identity no-op.
        for _f in ("disclosingEntityType", "receivingEntityType"):
            _v = full_data.get(_f)
            if isinstance(_v, str) and _v:
                full_data[_f] = _to_display_entity_type(_v, template_name, field_name=_f)
        # disputeForum is captured-not-rendered (see above): the clause keeps the default court
        # venue regardless of any supplied forum. That is the right call for material-term safety
        # (no silent court->arbitration flip), but a captured forum must NOT be silently dropped --
        # surface it as a non-fatal warning so the "captured for human review" promise is real
        # (Codex P2). Fires only when a caller explicitly supplied a forum, so normal NDA drafts
        # (no disputeForum) stay warning-free. Shared helper with draft() (T6).
        warnings = _forum_capture_warnings(
            full_data.get("disputeForum"), full_data.get("disputeForumRaw")
        )

        # P1.1: sanitize BEFORE hashing so audit_id -> rendered bytes stays 1:1.
        full_data = _sanitize_md_fields(full_data)
        data_hash = _compute_data_hash(
            full_data,
            overlay_map=overlay_map,
            clause_overrides=clause_overrides,
        )

        rendered = _render_unescaped(template_text, full_data, partials)

        return CiceroDraftResult(
            text=rendered,
            template_name=template_name,
            template_version=version,
            data_hash=data_hash,
            success=True,
            warnings=warnings,
        )
    except Exception as e:
        log.error(f"Cicero draft_with_data failed: {e}")
        return CiceroDraftResult(
            text="",
            template_name=template_name,
            template_version="0.0.0",
            data_hash="",
            success=False,
            error=str(e),
        )


# Reference .docx for pandoc --reference-doc styling
_REFERENCE_DOCX = Path(__file__).resolve().parent.parent / "data" / "templates" / "reference.docx"


def _find_pandoc() -> str | None:
    """Locate the pandoc binary. Returns path or None."""
    import shutil
    path = shutil.which("pandoc")
    if path:
        return path
    homebrew_path = "/opt/homebrew/bin/pandoc"
    if Path(homebrew_path).is_file():
        return homebrew_path
    return None


def markdown_to_docx(markdown_text: str, output_path: str) -> str:
    """Convert markdown to .docx via pandoc (preferred) or homebrew fallback.

    Pandoc produces professional output with proper heading styles, ordered
    lists for (a)(b)(c) sub-clauses, and styling from the reference template.
    Falls back to the homebrew python-docx converter if pandoc is unavailable.

    Returns the path to the generated file.
    """
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    pandoc = _find_pandoc()
    if pandoc:
        try:
            return _pandoc_to_docx(pandoc, markdown_text, str(out))
        except Exception as e:
            log.warning(f"Pandoc conversion failed, falling back to homebrew: {e}")

    return _homebrew_to_docx(markdown_text, str(out))


def _pandoc_to_docx(pandoc_path: str, markdown_text: str, output_path: str) -> str:
    """Convert markdown to .docx via pandoc subprocess."""
    import subprocess

    cmd = [pandoc_path, "-f", "markdown", "-t", "docx", "-o", output_path]
    if _REFERENCE_DOCX.is_file():
        cmd.extend(["--reference-doc", str(_REFERENCE_DOCX)])

    result = subprocess.run(
        cmd,
        input=markdown_text,
        capture_output=True,
        text=True,
        timeout=30,
    )
    if result.returncode != 0:
        raise RuntimeError(f"pandoc exited {result.returncode}: {result.stderr.strip()}")
    return output_path


def _homebrew_to_docx(markdown_text: str, output_path: str) -> str:
    """Fallback: convert markdown to .docx using python-docx (limited formatting)."""
    import re
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            i += 1
            continue
        if stripped.startswith("#"):
            level = 0
            while level < len(stripped) and stripped[level] == "#":
                level += 1
            if level < len(stripped) and stripped[level] == " ":
                text = stripped[level + 1:].strip()
                doc.add_heading(text, level=min(level, 4))
                i += 1
                continue
        para = doc.add_paragraph()
        _homebrew_render_inline(para, stripped)
        i += 1

    doc.save(output_path)
    return output_path


def _homebrew_render_inline(paragraph, text: str) -> None:
    """Render inline bold (**text**) into a python-docx paragraph."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
